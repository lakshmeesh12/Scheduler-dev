import os
import asyncio
from concurrent.futures import ThreadPoolExecutor
import fitz  # PyMuPDF
import docx
from PIL import Image
import io
import pytesseract
import re
import logging
from db.mongo.config import db as mongo_db
from langchain.text_splitter import RecursiveCharacterTextSplitter
import tiktoken
from typing import Dict, List, Tuple, Any, Optional, AsyncGenerator
import hashlib
from datetime import datetime, timezone
from utils.chatgpt import run_chatgpt
import zipfile
from uuid import uuid4
from utils.helper import compute_duration
import json
import time
from tenacity import retry, stop_after_attempt, wait_exponential
from bson import ObjectId

# Configure pytesseract path (update based on your system)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ChunkingPromptTemplate:
    def __init__(self, data: str):
        self.structure = f"""Extract resume details from the doctext below and return a Python-style JSON dictionary only.

** Extract contents based on titles and their datatypes mentioned below
    name: str,
    total_experience: int/None,
    email: str,
    linkedin_url: str/None,
    github_url: str/None,
    projects: list[dict(title, description, skills/tools, impact, start_date, end_date)],
    work_history: list[dict(company, designation, description, start_date, end_date)],
    primary_skills: dict(domain: dict(libraries: [], tools: [], concepts: [])),
    secondary_skills: dict(domain: dict(libraries: [], tools: [], concepts: [])),
    course: dict(institution, domain, level),
    certifications: list[str],
    education: dict(institution, degree, domain)

Input doctext:
{data}
"""
        self.skills_instructions = """
**Skills instructions and structure example**
{
  "Programming Languages": {
    "Languages": ["Python", "Java", "C++", "C", "C#", "JavaScript", "TypeScript", "Go", "Rust", "Kotlin", "Swift", "Ruby", "PHP", "R", "Scala", "Perl", "MATLAB", "SQL", "Bash"]
  },
  "Data Structures & Algorithms": {
    "Data Structures": ["Array", "Linked List", ...],
    "Algorithms": ["Binary Search", "Quick Sort", ...]
  },
  "Databases": {
    "Relational Databases": ["MySQL", "PostgreSQL", "Oracle", "SQL Server", "SQLite"],
    "NoSQL Databases": ["MongoDB", "Cassandra", "Redis", "CouchDB", "Neo4j", "Elasticsearch"],
    "Concepts": ["SQL", "ACID", "CAP Theorem", "Normalization", "Transactions", "Indexing"]
  },
  "Big Data": {
    "Frameworks": ["Hadoop", "Apache Spark", "Apache Flink", "Apache Kafka"],
    "Tools": ["Hive", "HBase", "Apache Pig"],
    "Concepts": ["Distributed Computing", "MapReduce", "ETL", "HDFS"]
  },
  "Machine Learning": {
    "Concepts": ["Supervised Learning", "Unsupervised Learning", "Reinforcement Learning", "Feature Engineering", "Model Evaluation", "Bias-Variance", "Cross-Validation", "Dimensionality Reduction", "Regularization", "Ensemble Methods", "Hyperparameter Tuning"],
    "Algorithms": ["Linear Regression", "Logistic Regression", "Decision Trees", "Random Forest", "SVM", "KNN", "K-Means", "Hierarchical Clustering", "Naive Bayes", "Gradient Boosting", "XGBoost", "LightGBM", "CatBoost", "AdaBoost"],
    "Libraries": ["scikit-learn", "Spark MLlib", "H2O.ai", "Weka"]
  }
}

** other instructions and constraints **

### Do not invent data that is not present.
### Expand short forms: e.g., "NLP" -> "Natural Language Processing".
### Return ONLY a JSON-compatible dictionary (no prose or code fences).
### The JSON must be valid and loadable by json.loads.
"""
        self.prompt = self.structure + self.skills_instructions


def _now_utc_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


class Resume:
    """
    Async resume pipeline: fully in-memory, concurrent, and with non-blocking MongoDB writes.
    """

    ALLOWED_EXTENSIONS = {"pdf", "docx"}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB per file
    MAX_FILES = 500

    # Budgets/limits for performance
    PDF_CONCURRENCY = 8
    OCR_CONCURRENCY = 2
    LLM_CONCURRENCY = 4
    TARGET_TOKEN_BUDGET = 6000  # keep prompts small for faster LLM latency

    def __init__(self, max_workers: Optional[int] = None):
        self.executor = ThreadPoolExecutor(max_workers=max_workers or max(os.cpu_count() * 2, 4))
        self.loop = asyncio.get_event_loop()

        # Separate semaphores: prevents CPU-bound and network-bound work from starving each other
        self.pdf_semaphore = asyncio.Semaphore(self.PDF_CONCURRENCY)
        self.ocr_semaphore = asyncio.Semaphore(self.OCR_CONCURRENCY)
        self.llm_semaphore = asyncio.Semaphore(self.LLM_CONCURRENCY)

        self.db = mongo_db
        self.collection = self.db["profiles"]

        self.tokenizer = tiktoken.encoding_for_model("text-embedding-3-small")
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=100,
            length_function=self._token_count,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
        )

        logger.info("Initialized async Resume pipeline")

    def _token_count(self, text: str) -> int:
        try:
            return len(self.tokenizer.encode(text))
        except Exception as e:
            logger.error(f"Error counting tokens: {str(e)}")
            return len(text)

    def validate_file(self, file) -> Dict[str, bool | str]:
        if not file.filename:
            return {"valid": False, "error": "File has no name"}
        file_ext = file.filename.split(".")[-1].lower()
        if file_ext not in self.ALLOWED_EXTENSIONS:
            return {"valid": False, "error": f"Invalid file type: {file_ext}. Allowed: {', '.join(self.ALLOWED_EXTENSIONS)}"}
        return {"valid": True}

    async def _image_to_text(self, image_data: bytes) -> str:
        try:
            async with self.ocr_semaphore:
                img = await self.loop.run_in_executor(self.executor, lambda: Image.open(io.BytesIO(image_data)).convert("RGB"))
                text = await self.loop.run_in_executor(self.executor, lambda: pytesseract.image_to_string(img, config="--psm 6"))
                return re.sub(r"\s+", " ", text).strip()
        except Exception as e:
            logger.warning(f"OCR failed: {e}")
            return ""

    async def get_text_docx(self, file_content: io.BytesIO) -> str:
        try:
            doc = await self.loop.run_in_executor(self.executor, lambda: docx.Document(file_content))
            pieces: List[str] = []

            for para in doc.paragraphs:
                txt = para.text.strip()
                if txt:
                    pieces.append(txt)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        txt = cell.text.strip()
                        if txt:
                            pieces.append(txt)

            # opportunistic OCR of embedded images (non-blocking and limited)
            ocr_tasks = []
            try:
                with zipfile.ZipFile(file_content) as zf:
                    for info in zf.infolist():
                        if info.filename.startswith("word/media/"):
                            data = zf.read(info)
                            ocr_tasks.append(self._image_to_text(data))
            except Exception:
                pass

            if ocr_tasks:
                # limit overall OCR concurrency
                ocr_results = await asyncio.gather(*ocr_tasks, return_exceptions=True)
                pieces.extend([t for t in ocr_results if isinstance(t, str) and t.strip()])

            return " ".join(pieces).strip()
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return ""

    async def _pdf_page_text(self, pdf, page_num: int) -> str:
        async with self.pdf_semaphore:
            try:
                page = await self.loop.run_in_executor(self.executor, lambda: pdf[page_num])
                text = await self.loop.run_in_executor(
                    self.executor,
                    lambda: page.get_text("text", flags=fitz.TEXTFLAGS_TEXT).replace("\n", " ").replace(" -", "-"),
                )
                text = re.sub(r"\s+", " ", text).strip()
                # heuristic: if page has almost no text, fallback to image OCR for that page
                if len(text) < 30:
                    pix = await self.loop.run_in_executor(self.executor, lambda: page.get_pixmap(dpi=200))
                    img_bytes = pix.tobytes("png")
                    ocr_text = await self._image_to_text(img_bytes)
                    return ocr_text or text
                return text
            except Exception as e:
                logger.warning(f"PDF page {page_num+1} failed: {e}")
                return ""

    async def get_text_pdf(self, file_content: io.BytesIO) -> str:
        try:
            with fitz.open(stream=file_content, filetype="pdf") as pdf:
                total_pages = len(pdf)
                tasks = [self._pdf_page_text(pdf, i) for i in range(total_pages)]
                out: List[str] = []
                # As pages finish, append to keep memory shallow
                for coro in asyncio.as_completed(tasks):
                    txt = await coro
                    if txt:
                        out.append(txt)
                return " ".join(out).strip()
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return ""

    async def extract_text(self, filename: str, file_content: io.BytesIO) -> Tuple[str, str]:
        ext = filename.split(".")[-1].lower()
        if ext == "pdf":
            return filename, await self.get_text_pdf(file_content)
        if ext == "docx":
            return filename, await self.get_text_docx(file_content)
        return filename, ""

    @retry(stop=stop_after_attempt(2), wait=wait_exponential(multiplier=1, min=2, max=8))
    async def _call_chatgpt_with_retry(self, prompt: str, system_message: str, temperature: float) -> str:
        try:
            async with self.llm_semaphore:
                resp = await run_chatgpt(prompt, system_message, temperature)
                if not resp or not isinstance(resp, str):
                    raise ValueError("Empty response from ChatGPT")
                return resp
        except Exception as e:
            logger.error(f"ChatGPT call failed: {e}")
            raise

    def _sanitize_skills(self, skills: Any) -> Dict[str, Dict[str, List[str]]]:
        sanitized: Dict[str, Dict[str, List[str]]] = {}
        if isinstance(skills, dict):
            for domain, categories in skills.items():
                if isinstance(categories, dict):
                    fixed = {}
                    for cat, items in categories.items():
                        if isinstance(items, list):
                            fixed[cat] = [str(s) for s in items if isinstance(s, (str, int, float)) and str(s).strip()]
                    if fixed:
                        sanitized[domain] = fixed
        return sanitized

    def _sanitize_json_response(self, response: str) -> str:
        try:
            cleaned = response.strip()
            if cleaned.startswith("```"):
                cleaned = cleaned.strip("`")
                cleaned = re.sub(r"^json", "", cleaned, flags=re.IGNORECASE).strip()
            cleaned = re.sub(r",\s*([}\]])", r"\1", cleaned)  # trailing commas
            cleaned = cleaned.replace("“", '"').replace("”", '"').replace("’", "'").replace("`", '"')
            # normalize single quotes for dicts that look Pythonic
            # only replace if it seems JSON-like with keys
            if '"' not in cleaned and re.search(r"\w+:", cleaned):
                cleaned = re.sub(r"'", '"', cleaned)
            return cleaned
        except Exception:
            return response

    def _truncate_to_budget(self, text: str, budget_tokens: int) -> str:
        toks = self.tokenizer.encode(text)
        if len(toks) <= budget_tokens:
            return text
        return self.tokenizer.decode(toks[:budget_tokens])

    async def _llm_extract(self, text: str) -> Dict[str, Any]:
        # keep prompt compact to reduce latency
        compact_text = self._truncate_to_budget(text, self.TARGET_TOKEN_BUDGET)
        prompt = ChunkingPromptTemplate(compact_text).prompt

        try:
            raw = await self._call_chatgpt_with_retry(
                prompt=prompt,
                system_message="You are an expert at extracting structured content from resume text. Only return valid JSON.",
                temperature=0.2,
            )
        except Exception:
            raw = ""

        cleaned = self._sanitize_json_response(raw) if raw else ""
        parsed: Dict[str, Any]
        if cleaned:
            try:
                parsed = json.loads(cleaned)
            except Exception as e:
                logger.warning(f"JSON parse failed, using fallback. Err: {e}")
                parsed = {}
        else:
            parsed = {}

        # fallback minimal shape to keep schema stable
        parsed = {
            "name": parsed.get("name") or None,
            "primary_skills": self._sanitize_skills(parsed.get("primary_skills", {})),
            "secondary_skills": self._sanitize_skills(parsed.get("secondary_skills", {})),
            "work_history": parsed.get("work_history", []) or [],
            "education": parsed.get("education", {}) or {},
            "email": parsed.get("email", "") or "",
            "linkedin_url": parsed.get("linkedin_url"),
            "github_url": parsed.get("github_url"),
            "projects": parsed.get("projects", []) or [],
            "course": parsed.get("course", {}) or {},
            "certifications": parsed.get("certifications", []) or [],
            "raw_text": text,
        }
        return parsed

    async def process_resume(self, filename: str, content: bytes) -> Optional[Dict[str, Any]]:
        try:
            file_content = io.BytesIO(content)
            _, text = await self.extract_text(filename, file_content)
            if not text:
                logger.warning(f"No text extracted from {filename}")
                return None

            parsed = await self._llm_extract(text)

            # Compute total experience
            total_exp = 0
            for rec in parsed.get("work_history", []):
                start = rec.get("start_date")
                end = rec.get("end_date") or "Sep 2025"
                if isinstance(end, str) and end.lower().strip() == "present":
                    end = "Sep 2025"
                total_exp += compute_duration(start, end)

            profile = {
                "profile_id": str(uuid4()),
                "file_name": filename,
                "file_hash": hashlib.sha256(content).hexdigest(),
                "total_experience": total_exp,
                "processed_at": _now_utc_iso(),
                "status": "COMPLETED",
                "active": True,
                **parsed,
            }

            # Non-blocking insert to Mongo
            try:
                insert_result = await self.collection.insert_one(profile)
                profile["_id"] = str(insert_result.inserted_id)
            except Exception as e:
                logger.warning(f"Mongo insert failed for {filename}: {e}")

            return profile
        except Exception as e:
            logger.error(f"Error processing {filename}: {e}")
            return None

    async def process_resumes_stream(
        self, files: List[Tuple[str, bytes]]
    ) -> AsyncGenerator[Dict[str, Any], None]:
        """
        Processes resumes concurrently and yields each result as soon as it's ready.
        Designed for streaming responses (NDJSON).
        """
        start = time.time()
        tasks = []
        for (fname, content) in files:
            tasks.append(asyncio.create_task(self.process_resume(fname, content)))

        success = 0
        failure = 0

        for task in asyncio.as_completed(tasks):
            result = await task
            if result:
                success += 1
                yield {"type": "item", "data": result}
            else:
                failure += 1
                # optionally include filename; we can’t easily recover it here from the task
                yield {"type": "error", "data": {"message": "Failed to process a file"}}

        elapsed = time.time() - start
        yield {
            "type": "summary",
            "data": {
                "message": "Parsing completed successfully!" if failure == 0 else "Parsing completed with errors!",
                "stats": {
                    "success_count": success,
                    "failure_count": failure,
                    "total_count": len(files),
                    "processing_time": elapsed,
                },
            },
        }

    def __del__(self):
        try:
            self.executor.shutdown(wait=False)
        except Exception as e:
            logger.error(f"Executor cleanup failed: {e}")
