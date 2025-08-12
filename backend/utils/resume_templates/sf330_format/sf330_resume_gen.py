import asyncio
import os
import json
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Pt  
from docx.oxml.ns import qn 
from copy import deepcopy


# Load state abbreviations from JSON file
def load_state_abbreviations():
    """Load state abbreviations from JSON file."""
    try:
        with open("state.json", "r", encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        print("Warning: state_abbreviations.json not found, using built-in mapping")
        return {
            "AL": "Alabama", "AK": "Alaska", "AZ": "Arizona", "AR": "Arkansas", "CA": "California", 
            "CO": "Colorado", "CT": "Connecticut", "DE": "Delaware", "FL": "Florida", "GA": "Georgia", 
            "HI": "Hawaii", "ID": "Idaho", "IL": "Illinois", "IN": "Indiana", "IA": "Iowa", 
            "KS": "Kansas", "KY": "Kentucky", "LA": "Louisiana", "ME": "Maine", "MD": "Maryland", 
            "MA": "Massachusetts", "MI": "Michigan", "MN": "Minnesota", "MS": "Mississippi", "MO": "Missouri", 
            "MT": "Montana", "NE": "Nebraska", "NV": "Nevada", "NH": "New Hampshire", "NJ": "New Jersey", 
            "NM": "New Mexico", "NY": "New York", "NC": "North Carolina", "ND": "North Dakota", "OH": "Ohio", 
            "OK": "Oklahoma", "OR": "Oregon", "PA": "Pennsylvania", "RI": "Rhode Island", "SC": "South Carolina", 
            "SD": "South Dakota", "TN": "Tennessee", "TX": "Texas", "UT": "Utah", "VT": "Vermont", 
            "VA": "Virginia", "WA": "Washington", "WV": "West Virginia", "WI": "Wisconsin", "WY": "Wyoming"
        }

STATE_ABBREVIATIONS = load_state_abbreviations()

def convert_state_to_abbreviation(state_text):
    """Convert state names to abbreviations."""
    if not state_text:
        return state_text
    
    state_text = state_text.strip()
    
    # If it's already an abbreviation (2 letters), return as is
    if len(state_text) == 2 and state_text.upper() in STATE_ABBREVIATIONS:
        return state_text.upper()
    
    # Convert full state name to abbreviation
    for abbrev, full_name in STATE_ABBREVIATIONS.items():
        if full_name.lower() in state_text.lower() or state_text.lower() in full_name.lower():
            return abbrev
    
    return state_text

def format_location_with_state_abbrev(location):
    """Format location to use state abbreviations."""
    if not location:
        return location
    
    parts = [part.strip() for part in location.split(',')]
    if len(parts) >= 2:
        city = parts[0]
        state_part = parts[1]
        state_abbrev = convert_state_to_abbreviation(state_part)
        return f"{city}, {state_abbrev}"
    
    return location

def calculate_experience_years(person_data):
    """
    Calculate total experience and years with current firm.
    
    Args:
        person_data (dict): Dictionary containing person information
        
    Returns:
        tuple: (total_years, current_firm_years) as strings
    """
    
    # Get total experience from input data
    total_years = str(person_data.get('total_experience', ''))
    
    # Calculate years with current firm from hire_date
    hire_date_str = person_data.get('hire_date', '')
    current_firm_years = ''
    
    if hire_date_str:
        try:
            # Parse the hire date string (format: "2013-05-19 18:30:50")
            hire_date = datetime.strptime(hire_date_str, "%Y-%m-%d %H:%M:%S")
            current_date = datetime.now()
            
            # Calculate years difference
            years_diff = current_date.year - hire_date.year
            
            # Adjust if we haven't reached the anniversary date this year
            if (current_date.month, current_date.day) < (hire_date.month, hire_date.day):
                years_diff -= 1
                
            current_firm_years = str(years_diff)
            
        except ValueError:
            try:
                # Try parsing just the date part if full datetime fails
                hire_date = datetime.strptime(hire_date_str.split()[0], "%Y-%m-%d")
                current_date = datetime.now()
                years_diff = current_date.year - hire_date.year
                if (current_date.month, current_date.day) < (hire_date.month, hire_date.day):
                    years_diff -= 1
                current_firm_years = str(years_diff)
            except ValueError:
                print(f"Warning: Could not parse hire date: {hire_date_str}")
                current_firm_years = ''
    
    return total_years, current_firm_years


async def resume_generation_sf330(person_data, template_path, output_path=None):
    try:
        """Main asynchronous function to process the SF330 document with all updates."""
        print("Person data provided as dictionary, processing directly!")
        
        loop = asyncio.get_event_loop()
        try:
            doc = await loop.run_in_executor(None, Document, template_path)
            print(f"Loaded document from {template_path}")
        except FileNotFoundError:
            print(f"Error: Template file not found at {template_path}")
            return
        except Exception as e:
            print(f"Error loading document: {str(e)}")
            return

        # Set default output path if not provided
        if not output_path:
            temp_dir = str(tempfile.gettempdir()).replace("\\",'/')
            output_path = os.path.join(temp_dir, f'converted_sf330_{asyncio.current_task().get_name()}.docx')
            print(f"No output path provided, using temporary path: {output_path}")
            
    except Exception as err:
        print(f"Error: {err}")

    
    # Step 2: Asynchronous helper function to set font style
    async def set_font_style(paragraph, font_name='Arial MT', font_size=Pt(7)):
        """Set font properties for all runs in a paragraph asynchronously."""
        loop = asyncio.get_event_loop()
        def sync_set_font_style():
            """Synchronous inner function to set font properties."""
            for run in paragraph.runs:
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
                run.font.size = font_size
        await loop.run_in_executor(None, sync_set_font_style)

    # Update personnel details
    async def update_personnel_details(doc, data):
        """Update personnel details (name, role, experience) in the document."""
        # Your JSON has separate first_name and last_name
        full_name = f"{data.get('name', '')}"
        
        # Use the new function to calculate experience years
        total_years, current_firm_years = data.get("total_experience"), data.get("years_experience_firm")
        
        for table in doc.tables:
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    cell_text = cell.text.strip().upper()
                    if "A. NAME" in cell_text:
                        row.cells[i].text = f"A. NAME\n{full_name}"
                        await set_font_style(row.cells[i].paragraphs[0])
                    elif "B. ROLE IN THIS CONTRACT" in cell_text:
                        role = data.get('job_title', '')
                        row.cells[i].text = f"B. ROLE IN THIS CONTRACT\n{role}"
                        await set_font_style(row.cells[i].paragraphs[0])
                    elif "1. TOTAL" in cell_text:
                        row.cells[i].text = f"1. TOTAL\n{total_years}"
                        await set_font_style(row.cells[i].paragraphs[0])
                    elif "2. WITH CURRENT FIRM" in cell_text:
                        row.cells[i].text = f"2. WITH CURRENT FIRM\n{current_firm_years}"
                        await set_font_style(row.cells[i].paragraphs[0])
        print("Step 3: Personnel details updated successfully!")

    async def update_firm_name_location(doc, person_data):
        """Update firm name and location in tables or paragraphs from work history."""
        found = False
       
        # Check work history for company and location
        work_history = person_data.get('work_history', [])
        if work_history and isinstance(work_history, list) and len(work_history) > 0:
            # Get the first record from work history
            first_record = work_history[0]
            company = first_record.get('company', '')
            location = first_record.get('location', '')
            if not location:
                location = person_data.get('location', '')
            formatted_location = format_location_with_state_abbrev(location)
            firm_name_location = f"{company}, {formatted_location}"
        else:
            # Fallback if no work history is available
            location = person_data.get('location', '')
            formatted_location = format_location_with_state_abbrev(location)
            firm_name_location = formatted_location
     
        # Update the document with firm name and location
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if "FIRM NAME" in row_text.upper() and "LOCATION" in row_text.upper():
                    for cell in row.cells:
                        if "FIRM NAME" in cell.text.upper() and "LOCATION" in cell.text.upper():
                            cell.text = f"D. FIRM NAME AND LOCATION\n{firm_name_location}"
                            await set_font_style(cell.paragraphs[0])
                            found = True
                            break
                    if found:
                        break
            if found:
                break
       
        if not found:
            for paragraph in doc.paragraphs:
                if "FIRM NAME" in paragraph.text.upper() and "LOCATION" in paragraph.text.upper():
                    paragraph.text = f"D. FIRM NAME AND LOCATION\n{firm_name_location}"
                    await set_font_style(paragraph)
                    break
       
        print("Step 4: Firm name and location updated successfully!")

    # Update education and registration (corrected for certifications)
    async def update_education_and_registration(doc, person_data):
        """Update education and professional registration details with certifications."""
        found_education = False
        found_registration = False
        
        # Format education - simplified to just show degree and specialization
        education_list = person_data.get('education', [])
        if education_list:
            first_education = education_list[0]
            degree = first_education.get('degree', '')
            specialization = first_education.get('major_or_specialization', '')
            education_text = f"{degree}, {specialization}" if degree and specialization else f"{degree}{specialization}"
        else:
            education_text = ''
        
        # Format certifications with state abbreviations
        certifications_list = person_data.get('certifications', [])
        if certifications_list:
            formatted_certs = []
            for cert in certifications_list:
                formatted_cert = cert
                for abbrev, full_name in STATE_ABBREVIATIONS.items():
                    if full_name in cert:
                        formatted_cert = cert.replace(full_name, abbrev)
                        break
                formatted_certs.append(formatted_cert)
            registration_text = "\n".join(formatted_certs)
        else:
            registration_text = "No certifications provided"

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip().upper()
                    if not found_education and "E. EDUCATION" in cell_text:
                        cell.text = f"E. EDUCATION (Degree and Specialization)\n{education_text}"
                        await set_font_style(cell.paragraphs[0])
                        found_education = True
                    # Check for both correct and typo versions
                    if not found_registration and ("F. CURRENT PROFESSIONAL REGISTRATION" in cell_text or "F. CUREENT PROFESSIONAL REGISTRATION" in cell_text):
                        cell.text = f"F. CURRENT PROFESSIONAL REGISTRATION (State and Discipline)\n{registration_text}"
                        await set_font_style(cell.paragraphs[0])
                        found_registration = True
                    if found_education and found_registration:
                        break
                if found_education and found_registration:
                    break
        if not found_education:
            print("Warning: 'E. EDUCATION' section not found in the document.")
        if not found_registration:
            print("Warning: 'F. CURRENT PROFESSIONAL REGISTRATION' section not found in the document.")
        print("Step 5: Education and certifications updated successfully!")

    # Update other professional qualifications
    async def update_other_professional_qualifications(doc, person_data):
        """Update other professional qualifications with bio sketch description."""
        found = False
        # Your bio_sketch is nested in an array
        bio_sketch_list = person_data.get("bio_sketch", [])
        if bio_sketch_list and isinstance(bio_sketch_list, list) and len(bio_sketch_list) > 0:
            bio_sketch = bio_sketch_list[0].get("bio_sketch", "")
        else:
            bio_sketch = ""
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "OTHER PROFESSIONAL QUALIFICATIONS" in cell.text.upper():
                        cell.text = f"G. OTHER PROFESSIONAL QUALIFICATIONS (Publications, Organizations, Training, Awards, etc.)\n{bio_sketch}"
                        await set_font_style(cell.paragraphs[0])
                        found = True
                        break
                if found:
                    break
            if found:
                break
        if not found:
            for paragraph in doc.paragraphs:
                if "OTHER PROFESSIONAL QUALIFICATIONS" in paragraph.text.upper():
                    paragraph.text = f"G. OTHER PROFESSIONAL QUALIFICATIONS (Publications, Organizations, Training, Awards, etc.)\n{bio_sketch}"
                    await set_font_style(paragraph)
                    break
        print("Step 6: Other professional qualifications updated successfully!")

    # Update relevant projects
    async def update_relevant_projects(doc, person_data):
        """Update relevant projects table with titles on new lines, limited to top 5 projects."""
        projects_table = None
        for table in doc.tables:
            for row in table.rows:
                if "(1) TITLE AND LOCATION" in row.cells[1].text.upper():
                    projects_table = table
                    break
            if projects_table: 
                break 

        if projects_table:
            template_block = [deepcopy(row._element) for row in projects_table.rows[:3]]
            while len(projects_table.rows) > 0:
                projects_table._tbl.remove(projects_table.rows[0]._tr)
            
            projects = person_data.get("projects", [])[:5]
            for proj_idx in range(len(projects)):
                for row_element in template_block:
                    projects_table._tbl.append(deepcopy(row_element))

            for proj_idx, project in enumerate(projects):
                base_row = proj_idx * 3 
                label = f"{proj_idx + 1}." 
                start_year = project.get('start_date', '')
                end_year = project.get('end_date', '')
                description = project.get('description', 'No description provided')
                
                # Format location with state abbreviation
                original_location = project.get('location', '')
                formatted_location = format_location_with_state_abbrev(original_location)

                professional_services_date = project.get('professional_services_completion_date', '') 
                construction_completion_date = project.get('construction_completion_date', '') 
                project_design_costs = f"$ {project.get('project_design_costs', '')}".rstrip("$ ")
                project_construction_costs = f"$ {project.get('project_construction_costs', '')}".rstrip("$ ")

                row1 = projects_table.rows[base_row] 
                row1.cells[0].text = label 
                row1.cells[1].text = f"(1) TITLE AND LOCATION (City and State)\n{project.get('title', '')}\n{formatted_location}"
                year_completed = construction_completion_date if construction_completion_date else end_year 
                if year_completed and year_completed != 'null' and year_completed.strip(): 
                    row1.cells[2].text = f"(2) YEAR COMPLETED\n\n{year_completed}" 
                else:
                    row1.cells[2].text = ""  
                for cell in row1.cells: 
                    await set_font_style(cell.paragraphs[0]) 

                row2 = projects_table.rows[base_row + 1] 
                row2.cells[0].text = label 
                row2.cells[1].text = f"(1) TITLE AND LOCATION (City and State)\n{project.get('title', 'N/A')}\n{formatted_location}"

                if professional_services_date and professional_services_date != 'null' and professional_services_date.strip():
                    prof_services_text = f"PROFESSIONAL SERVICES\n{professional_services_date}"
                elif start_year and start_year != 'null' and start_year.strip():
                    prof_services_text = f"PROFESSIONAL SERVICES\n{start_year}"
                else:
                    prof_services_text = ""  
                row2.cells[2].text = prof_services_text 
            
                if construction_completion_date and construction_completion_date != 'null' and construction_completion_date.strip():
                    construction_text = f"CONSTRUCTION (If applicable)\n{construction_completion_date}"
                elif end_year and end_year != 'null' and end_year.strip():
                    construction_text = f"CONSTRUCTION (If applicable)\n{end_year}"
                else:
                    construction_text = "" 
            
                if len(row2.cells) > 3: 
                    row2.cells[3].text = construction_text 
            
                for cell in row2.cells:
                    await set_font_style(cell.paragraphs[0])

                row3 = projects_table.rows[base_row + 2]
                row3.cells[0].text = label

                enhanced_description = description 
            
                cost_details = []
                if project_design_costs:
                    cost_details.append(f"Cost (fee): {project_design_costs}")
                if project_construction_costs:
                    cost_details.append(f"Cost (construction): {project_construction_costs}") 
            
                if cost_details:
                    enhanced_description += "\n\n" + " | ".join(cost_details)
            
                row3.cells[1].text = f"(3) BRIEF DESCRIPTION (Brief scope, size, cost, etc.) AND SPECIFIC ROLE\n{enhanced_description}"
                
                # Add checkbox logic for "performed with current firm"
                performed_with_current_firm = project.get('isPerformedWithCurrentFirm', False)
                if performed_with_current_firm:
                    checkbox_text = "☑ Check if project performed with current firm"
                else:
                    checkbox_text = "☐ Check if project performed with current firm"
                row3.cells[2].text = checkbox_text
                for cell in row3.cells:
                    await set_font_style(cell.paragraphs[0])
            print("Step 7: Relevant projects updated successfully!")

    # Step 8: Call all update functions concurrently
    await asyncio.gather(
        update_personnel_details(doc, person_data),
        update_firm_name_location(doc, person_data),
        update_education_and_registration(doc, person_data),
        update_other_professional_qualifications(doc, person_data),
        update_relevant_projects(doc, person_data)
    )

    # Step 9: Save the updated document 
    await loop.run_in_executor(None, doc.save, output_path)
    print(f"Step 9: Updated document saved as {output_path}")
    return output_path 
    
async def main():
    """Load JSON manually and generate SF330 resume"""
    
    # Load JSON file manually
    try:
        with open("input.json", "r", encoding='utf-8') as f:
            person_data = json.load(f)
        print("Successfully loaded JSON data")
        print(f"Processing resume for: {person_data.get('first_name', '')} {person_data.get('last_name', '')}")
    except FileNotFoundError:
        print("Error: input.json file not found")
        return
    except json.JSONDecodeError as e:
        print(f"Error: Invalid JSON format. {str(e)}")
        return
    except Exception as e:
        print(f"Error loading JSON: {str(e)}")
        return
    
    template_path = "Template form sf330.docx" 
    output_path = "sf330_from_json.docx"
    
    result = await resume_generation_sf330(
        person_data=person_data,  # Pass the loaded data
        template_path=template_path,
        output_path=output_path
    )
    
    if result: 
        print(f"SF330 resume generated successfully: {result}")
    else:
        print("Failed to generate SF330 resume")


if __name__ == "__main__":
    # Run the main function
    asyncio.run(main())