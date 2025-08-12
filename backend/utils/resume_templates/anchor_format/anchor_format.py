import json
import asyncio
import os
import tempfile
import aiofiles
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

async def load_config(config_path):
    try:
        async with aiofiles.open(config_path, 'r', encoding="utf-8") as f:
            config = json.loads(await f.read())
        config["colors"] = {k: RGBColor(*v) for k, v in config["colors"].items()}
        config["defaults"]["font_size"] = config["defaults"]["font_size"]
        config["defaults"]["projects_font_size"] = Pt(config["defaults"]["projects_font_size"])
        config["defaults"]["separator_font_size"] = Pt(config["defaults"]["separator_font_size"])
        config['formatting']['projects_separator'] = "\n–––––––––––––––––––––––––––––––––––––––––––––––––––"
        return config
    except Exception as err:
        print("Error in load_config: ",err)

async def initialize_config():
    try:
        global CONFIG
        CURR_DIR = os.getcwd().replace("\\","/")
        CURR_DIR += "/utils/resume_templates/anchor_format/"
        print("stage 1 ",CURR_DIR)
        CONFIG = await load_config(CURR_DIR+'config.json')
        print(CONFIG)
    except Exception as err:
        print("Error in initialize_config: ", err)

async def resume_generation_anchor(resume_data, template_path, output_path):
    try:
        print("Resume data provided as dictionary, processing directly!")
        
        loop = asyncio.get_event_loop()
        doc = await loop.run_in_executor(None, Document, template_path)

        await replace_placeholders(doc, resume_data)
        await replace_summary(doc, resume_data)
        await replace_education(doc, resume_data.get(CONFIG["json_keys"]["education"], []))
        await replace_certifications(doc, resume_data.get(CONFIG["json_keys"]["certifications"], []))
        await replace_work_history(doc, resume_data.get("work_history", []))  # Moved here
        await replace_projects(doc, resume_data.get(CONFIG["json_keys"]["projects"], []))
        await replace_publications(doc, resume_data.get(CONFIG["json_keys"]["publications"], []))
        await replace_awards(doc, resume_data.get(CONFIG["json_keys"]["awards"], []))

        await loop.run_in_executor(None, doc.save, output_path)
        print(f"Resume successfully updated and saved as {output_path}!")
        return output_path
    except Exception as err:
        print(f"Error in resume_generation : {err}")
        return ""

async def replace_placeholders(doc, resume_data):
    name_value = resume_data.get("name", "")
    title_value = resume_data.get("job_title", "")
    placeholders = {
        CONFIG["placeholders"]["name"]: name_value,
        CONFIG["placeholders"]["title"]: title_value
    }
    
    loop = asyncio.get_event_loop()
    
    for para in doc.paragraphs:
        for placeholder, replacement in placeholders.items():
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, replacement)
                for run in para.runs:
                    run.font.name = CONFIG["defaults"]["font_name"]
                    run.font.size = Pt(30)
    
    for section in doc.sections:
        for para in section.header.paragraphs:
            for placeholder, replacement in placeholders.items():
                if placeholder in para.text:
                    para.text = para.text.replace(placeholder, replacement)
                    for run in para.runs:
                        run.font.name = CONFIG["defaults"]["font_name"]
                        run.font.size = Pt(12)

async def replace_summary(doc, resume_data):
    
    summary_text = resume_data.get("bio_sketch", "")
    summary_replaced = False
    
    for i, para in enumerate(doc.paragraphs):
        full_text = "".join(run.text for run in para.runs)
        if CONFIG["placeholders"]["summary"] in full_text:
            para.clear()
            run = para.add_run(summary_text)
            run.font.name = CONFIG["defaults"]["font_name"]
            run.font.size = Pt(11)
            summary_replaced = True
            
            if i + 1 < len(doc.paragraphs):
                next_text = "".join(run.text for run in doc.paragraphs[i + 1].runs)
                if CONFIG["placeholders"]["summary_continuation"] in next_text:
                    doc.paragraphs[i + 1].clear()
            break
    
    if not summary_replaced:
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        full_text = "".join(run.text for run in para.runs)
                        print(f"Table {table_idx}, Row {row_idx}, Cell {cell_idx}, Para {para_idx}: '{full_text}'")
                        if CONFIG["placeholders"]["summary"] in full_text:
                            print(f"Summary placeholder found in table {table_idx}, row {row_idx}, cell {cell_idx}, para {para_idx}: '{full_text}'")
                            para.clear()
                            run = para.add_run(summary_text)
                            run.font.name = CONFIG["defaults"]["font_name"]
                            run.font.size = Pt(11)
                            summary_replaced = True
                            
                            if para_idx + 1 < len(cell.paragraphs):
                                next_text = "".join(run.text for run in cell.paragraphs[para_idx + 1].runs)
                                if CONFIG["placeholders"]["summary_continuation"] in next_text:
                                    print(f"Clearing continuation paragraph in same cell: '{next_text}'")
                                    cell.paragraphs[para_idx + 1].clear()
                            break
                    if summary_replaced:
                        break
                if summary_replaced:
                    break
            if summary_replaced:
                break
    
    if not summary_replaced:
        new_para = doc.add_paragraph()
        run = new_para.add_run("")
        run.font.name = CONFIG["defaults"]["font_name"]
        run.font.size = Pt(11)

async def replace_education(doc, education):
    if not education:
        education_text = ""
    else:
        education_text = "\n".join([
            CONFIG["formatting"]["education_format"].format(
                Degree=edu["degree"],
                Major=edu["major_or_specialization"],
                Institution=edu["college_institution"],
                YearOfGraduation=""
            ).rstrip(", ") for edu in education
        ])
    
    replaced = False
    placeholder_texts = CONFIG["placeholders"]["education"]
    
    for para in doc.paragraphs:
        if any(placeholder in para.text for placeholder in placeholder_texts):
            if not replaced:
                para.clear()
                run = para.add_run(education_text)
                run.font.name = CONFIG["defaults"]["font_name"]
                run.font.size = Pt(11)
                replaced = True
            else:
                para.clear()
    
    if not replaced:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if any(placeholder in para.text for placeholder in placeholder_texts):
                            if not replaced:
                                para.clear()
                                run = para.add_run(education_text)
                                run.font.name = CONFIG["defaults"]["font_name"]
                                run.font.size = Pt(11)
                                replaced = True
                            else:
                                para.clear()
    
    if not replaced:
        new_para = doc.add_paragraph()
        run = new_para.add_run(education_text)
        run.font.name = CONFIG["defaults"]["font_name"]
        run.font.size = Pt(11)

async def replace_certifications(doc, certifications):
    cert_text = certifications if certifications else [""]
    cert_font_name = "Myriad Pro"
    cert_font_size = Pt(9)
    cert_section_found = False
    placeholder_text = CONFIG["placeholders"]["certifications"]
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs_to_remove = []
                for j, para in enumerate(cell.paragraphs):
                    full_text = "".join(run.text for run in para.runs).strip()
                    if placeholder_text in full_text and not cert_section_found:
                        cert_section_found = True
                        para.clear()
                        run = para.add_run(placeholder_text)
                        run.bold = True
                        run.font.name = cert_font_name
                        run.font.size = cert_font_size
                    elif "Discipline or Type" in full_text or CONFIG["placeholders"]["certifications_detail"] in full_text:
                        paragraphs_to_remove.append(j)
                    elif CONFIG["placeholders"]["certifications_work_history"] in full_text or full_text.startswith(CONFIG["formatting"]["bullet_symbol"].strip()):
                        paragraphs_to_remove.append(j)
                for j in sorted(paragraphs_to_remove, reverse=True):
                    cell.paragraphs[j]._element.getparent().remove(cell.paragraphs[j]._element)
                if cert_section_found:
                    for cert in cert_text:
                        new_para = cell.add_paragraph()
                        run = new_para.add_run(cert)
                        run.font.name = cert_font_name
                        run.font.size = cert_font_size
                        new_para.style.font.name = cert_font_name
                        new_para.style.font.size = cert_font_size
                        new_para.paragraph_format.left_indent = CONFIG["formatting"]["cert_indent"]
                    return
    
    if not cert_section_found:
        paragraphs_to_remove = []
        for i, para in enumerate(doc.paragraphs):
            full_text = "".join(run.text for run in para.runs).strip()
            if placeholder_text in full_text:
                cert_section_found = True
                para.clear()
                run = para.add_run(placeholder_text)
                run.bold = True
                run.font.name = cert_font_name
                run.font.size = cert_font_size
            elif CONFIG["placeholders"]["certifications_detail"] in full_text or CONFIG["placeholders"]["certifications_work_history"] in full_text or full_text.startswith(CONFIG["formatting"]["bullet_symbol"].strip()):
                paragraphs_to_remove.append(i)
        for i in sorted(paragraphs_to_remove, reverse=True):
            doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
        
        if cert_section_found:
            for i, para in enumerate(doc.paragraphs):
                if placeholder_text in "".join(run.text for run in para.runs):
                    for cert in cert_text:
                        new_para = doc.add_paragraph()
                        run = new_para.add_run(cert)
                        run.font.name = cert_font_name
                        run.font.size = cert_font_size
                        new_para.style.font.name = cert_font_name
                        new_para.style.font.size = cert_font_size
                    break
    
    if not cert_section_found:
        new_para = doc.add_paragraph()
        run = new_para.add_run(placeholder_text)
        run.bold = True
        run.font.name = cert_font_name
        run.font.size = cert_font_size
        for cert in cert_text:
            cert_para = doc.add_paragraph()
            run = cert_para.add_run(cert)
            run.font.name = cert_font_name
            run.font.size = cert_font_size

async def replace_work_history(doc, work_history):
    if not work_history:
        work_entries = [{"title": "", "description": ""}]
    else:
        work_entries = []
        for entry in work_history:
            end_year_str = entry.get('end_year') or 'Present'
            if isinstance(end_year_str, (int, float)): end_year_str = str(int(end_year_str))
            elif isinstance(end_year_str, str) and end_year_str.lower() == 'present': end_year_str = 'Present'
            title_line = f"{entry.get('company', 'N/A')}, {entry.get('designation', 'N/A')}, {entry.get('start_year', 'N/A')} to {end_year_str}"
            work_entries.append({"title": title_line, "description": entry.get("description", "")})

    cert_heading_text = CONFIG["placeholders"]["certifications"]
    found_cell = None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cert_heading_text in cell.text:
                    for p in cell.paragraphs:
                        para_text_norm = ' '.join(p.text.split())
                        if para_text_norm == cert_heading_text and any(r.bold for r in p.runs if r.text.strip()):
                            found_cell = cell
                            break
                if found_cell: break
            if found_cell: break
        if found_cell: break

    if found_cell:
        heading_font_name = "Myriad Pro"
        heading_font_size = Pt(9)
        details_font_name = "Myriad Pro"
        details_font_size = Pt(9)
        dark_teal_color = CONFIG["colors"].get("dark_teal", RGBColor(0, 0, 0))
        details_indent_emu = int(CONFIG["formatting"].get("cert_indent", 228600))

        heading_para = found_cell.add_paragraph()
        heading_run = heading_para.add_run("Work History")
        heading_run.font.name = heading_font_name
        heading_run.font.size = heading_font_size
        heading_run.bold = True
        heading_run.font.color.rgb = dark_teal_color
        heading_para.paragraph_format.left_indent = details_indent_emu
        heading_para.paragraph_format.space_before = Pt(10)
        heading_para.paragraph_format.space_after = Pt(2)

        for i, entry in enumerate(work_entries):
            title_para = found_cell.add_paragraph()
            title_run = title_para.add_run(entry["title"])
            title_run.font.name = details_font_name
            title_run.font.size = details_font_size
            title_para.paragraph_format.left_indent = details_indent_emu
            title_para.paragraph_format.space_before = Pt(0)
            title_para.paragraph_format.space_after = Pt(0)

            description = entry.get("description", "")
            if description and description.strip():
                desc_para = found_cell.add_paragraph()
                desc_run = desc_para.add_run(description)
                desc_run.font.name = details_font_name
                desc_run.font.size = details_font_size
                desc_para.paragraph_format.left_indent = details_indent_emu
                desc_para.paragraph_format.space_before = Pt(0)
                desc_para.paragraph_format.space_after = Pt(6)
            else:
                title_para.paragraph_format.space_after = Pt(6)
    else:
        temp_doc = Document()
        wh_elements = []
        heading_para = temp_doc.add_paragraph("Work History")
        run = heading_para.runs[0]
        run.bold = True
        run.font.size = Pt(11)
        heading_para.paragraph_format.space_before = Pt(12)
        heading_para.paragraph_format.space_after = Pt(6)
        wh_elements.append(heading_para._element)

        for entry in work_entries:
            title_para = temp_doc.add_paragraph(entry["title"])
            title_para.runs[0].font.name = CONFIG["defaults"]["font_name"]
            title_para.runs[0].font.size = Pt(10)
            title_para.paragraph_format.space_before = Pt(0)
            title_para.paragraph_format.space_after = Pt(0)
            wh_elements.append(title_para._element)

            description = entry.get("description", "")
            if description and description.strip():
                desc_para = temp_doc.add_paragraph(description)
                desc_para.runs[0].font.name = CONFIG["defaults"]["font_name"]
                desc_para.runs[0].font.size = Pt(10)
                desc_para.paragraph_format.left_indent = Inches(0.25)
                desc_para.paragraph_format.space_before = Pt(0)
                desc_para.paragraph_format.space_after = Pt(6)
                wh_elements.append(desc_para._element)
            else:
                title_para.paragraph_format.space_after = Pt(6)

        for element in wh_elements:
            doc.element.body.append(element)

async def get_table_cell_format(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.font and run.font.name and run.font.size and run.font.color and run.font.color.rgb:
                        return run.font.name, run.font.size, run.font.color.rgb
    return None, None, None

async def replace_projects(doc, projects):
    placeholders = CONFIG["placeholders"]["projects"]
    project_table = None
    project_rows = []
    
    for table in doc.tables:
        placeholder_rows = []
        for i, row in enumerate(table.rows):
            row_has_placeholders = any(
                any(placeholder in para.text for para in cell.paragraphs)
                for placeholder in placeholders
                for cell in row.cells
            )
            if row_has_placeholders:
                if project_table is None:
                    project_table = table
                placeholder_rows.append(i)
        
        if project_table == table and placeholder_rows:
            project_rows = placeholder_rows
            break
    
    if not project_table:
        return
    
    if not projects:
        if project_rows:
            first_row = project_table.rows[project_rows[0]]
            for cell in first_row.cells:
                for para in cell.paragraphs:
                    para.clear()
            
            if len(first_row.cells) > 1:
                no_projects_para = first_row.cells[1].add_paragraph()
                run = no_projects_para.add_run("")
                run.font.name = CONFIG["defaults"]["projects_font_name"]
                run.font.size = CONFIG["defaults"]["projects_font_size"]
                run.font.bold = True
            
            for i in sorted(project_rows[1:], reverse=True):
                project_table._element.remove(project_table.rows[i]._element)
        return
    
    font_name, font_size, _ = await get_table_cell_format(project_table)
    font_name = font_name or CONFIG["defaults"]["projects_font_name"]
    font_size = font_size or CONFIG["defaults"]["projects_font_size"]
    
    used_rows = []
    
    for i, proj_idx in enumerate(project_rows):
        if i >= len(projects):
            break
            
        row = project_table.rows[proj_idx]
        used_rows.append(proj_idx)
        
        for cell in row.cells:
            for para in cell.paragraphs[1:]:
                para._element.getparent().remove(para._element)
            if cell.paragraphs:
                cell.paragraphs[0].clear()
        
        if len(row.cells) >= 2:
            title_para = row.cells[0].paragraphs[0] if row.cells[0].paragraphs else row.cells[0].add_paragraph()
            title_run = title_para.add_run(projects[i]["title"])
            title_run.font.name = font_name
            title_run.font.size = font_size
            title_run.font.color.rgb = CONFIG["colors"]["dark_teal"]
            title_run.bold = True
            
            client_para = row.cells[0].add_paragraph()
            client_run = client_para.add_run(projects[i]["institute"])
            client_run.font.name = font_name
            client_run.font.size = font_size
            client_run.font.color.rgb = CONFIG["colors"]["teal_accent_2"]
            
            location_para = row.cells[0].add_paragraph()
            location_run = location_para.add_run(projects[i]["location"])
            location_run.font.name = font_name
            location_run.font.size = font_size
            location_run.font.color.rgb = CONFIG["colors"]["teal_accent_2"]
            
            desc_para = row.cells[1].paragraphs[0] if row.cells[1].paragraphs else row.cells[1].add_paragraph()
            desc_run = desc_para.add_run(projects[i]["description"])
            desc_run.font.name = font_name
            desc_run.font.size = font_size
            
            row.cells[1].width = Inches(4.0)
            
            if i < len(projects) - 1:
                separator_para = row.cells[1].add_paragraph()
                pPr = separator_para._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '4')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), '009999')
                pBdr.append(bottom)
                pPr.append(pBdr)
                
                separator_para.paragraph_format.space_before = Pt(8)
                separator_para.paragraph_format.space_after = Pt(8)
    
    unused_rows = [idx for idx in project_rows if idx not in used_rows]
    for i in sorted(unused_rows, reverse=True):
        project_table._element.remove(project_table.rows[i]._element)
    
    if len(projects) > len(project_rows):
        for i in range(len(project_rows), len(projects)):
            new_row = project_table.add_row()
            
            title_para = new_row.cells[0].add_paragraph()
            title_run = title_para.add_run(projects[i]["title"])
            title_run.font.name = font_name
            title_run.font.size = font_size
            title_run.font.color.rgb = CONFIG["colors"]["dark_teal"]
            title_run.bold = True
            
            client_para = new_row.cells[0].add_paragraph()
            client_run = client_para.add_run(projects[i]["institute"])
            client_run.font.name = font_name
            client_run.font.size = font_size
            client_run.font.color.rgb = CONFIG["colors"]["teal_accent_2"]
            
            location_para = new_row.cells[0].add_paragraph()
            location_run = location_para.add_run(projects[i]["location"])
            location_run.font.name = font_name
            location_run.font.size = font_size
            location_run.font.color.rgb = CONFIG["colors"]["teal_accent_2"]
            
            desc_para = new_row.cells[1].add_paragraph()
            desc_run = desc_para.add_run(projects[i]["description"])
            desc_run.font.name = font_name
            desc_run.font.size = font_size
            
            new_row.cells[1].width = Inches(4.0)
            
            if i < len(projects) - 1:
                separator_para = new_row.cells[1].add_paragraph()
                pPr = separator_para._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '4')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), '009999')
                pBdr.append(bottom)
                pPr.append(pBdr)
                
                separator_para.paragraph_format.space_before = Pt(8)
                separator_para.paragraph_format.space_after = Pt(8)

async def format_bullet_paragraph(para, text):
    para.clear()
    para.add_run(f"{CONFIG['formatting']['bullet_symbol']}{text}")
    para.paragraph_format.left_indent = CONFIG["formatting"]["bullet_left_indent"]
    para.paragraph_format.first_line_indent = CONFIG["formatting"]["bullet_first_line_indent"]

async def replace_publications(doc, publications):
    pub_text = publications if publications else [""]
    placeholder_text = CONFIG["placeholders"]["publications"]
    publications_heading_index = None
    placeholder_index = None
    
    for i, para in enumerate(doc.paragraphs):
        full_text = "".join(run.text for run in para.runs)
        if "Publications" in full_text:
            publications_heading_index = i
        if placeholder_text in full_text:
            placeholder_index = i
            break
    
    target_index = placeholder_index if placeholder_index is not None else (
        publications_heading_index + 1 if publications_heading_index is not None else None
    )
    
    if target_index is not None:
        if target_index < len(doc.paragraphs):
            doc.paragraphs[target_index].clear()
        else:
            doc.add_paragraph()
        
        font_name = CONFIG["defaults"]["font_name"]
        font_size = Pt(11)
        
        first_bullet_para = doc.paragraphs[target_index]
        await format_bullet_paragraph(first_bullet_para, pub_text[0])
        first_bullet_para.runs[0].font.name = font_name
        first_bullet_para.runs[0].font.size = font_size
        
        for pub in pub_text[1:]:
            bullet_para = doc.add_paragraph()
            await format_bullet_paragraph(bullet_para, pub)
            bullet_para.runs[0].font.name = font_name
            bullet_para.runs[0].font.size = font_size
            first_bullet_para._element.addnext(bullet_para._element)
        
        for para in doc.paragraphs:
            if placeholder_text in "".join(run.text for run in para.runs) and para != doc.paragraphs[target_index]:
                para.clear()
    else:
        doc.add_paragraph("Publications").style = "Heading 1"
        font_name = CONFIG["defaults"]["font_name"]
        font_size = Pt(11)
        
        first_bullet_para = doc.add_paragraph()
        await format_bullet_paragraph(first_bullet_para, pub_text[0])
        first_bullet_para.runs[0].font.name = font_name
        first_bullet_para.runs[0].font.size = font_size
        
        for pub in pub_text[1:]:
            bullet_para = doc.add_paragraph()
            await format_bullet_paragraph(bullet_para, pub)
            bullet_para.runs[0].font.name = font_name
            bullet_para.runs[0].font.size = font_size
            first_bullet_para._element.addnext(bullet_para._element)

async def format_paragraph(para, text, numbered=True):
    para.clear()
    run = para.add_run(text)
    run.font.name = CONFIG["defaults"]["font_name"]
    run.font.size = Pt(11)
    pPr = para._element.get_or_add_pPr()
    if numbered:
        numPr = OxmlElement('w:numPr')
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), CONFIG["formatting"]["numbering_id"])
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), CONFIG["formatting"]["numbering_level"])
        numPr.append(ilvl)
        numPr.append(numId)
        pPr.append(numPr)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), CONFIG["formatting"]["spacing_after"])
    pPr.append(spacing)

async def replace_awards(doc, awards):
    awards_list = awards if awards and not (isinstance(awards, dict) and awards.get("$numberDouble") == "NaN") else [""]
    awards_section_found = False
    placeholder_paragraphs = []
    awards_heading_index = None
    
    for i, para in enumerate(doc.paragraphs):
        para_text = "".join(run.text for run in para.runs).strip()
        if "Awards and Honours" in para_text or "Awards and Honors" in para_text:
            awards_heading_index = i
            awards_section_found = True
            continue
        if awards_section_found:
            if ("Award Name" in para_text or 
                CONFIG["placeholders"]["awards"] in para_text or
                para_text.startswith("*") or 
                para_text.startswith("•") or
                para_text.startswith("1.")):
                placeholder_paragraphs.append(i)
            elif (para_text and para_text[0].isupper() and 
                  not para_text.startswith("1.") and not para_text.startswith("•")):
                break
    
    if awards_section_found:
        for idx in sorted(placeholder_paragraphs, reverse=True):
            if idx < len(doc.paragraphs):
                para = doc.paragraphs[idx]
                para._element.getparent().remove(para._element)
        
        insertion_point = awards_heading_index + 1
        if awards_list:
            first_award_para = doc.add_paragraph()
            await format_paragraph(first_award_para, awards_list[0], numbered=True)
            doc.paragraphs[awards_heading_index]._element.addnext(first_award_para._element)
            
            for award in awards_list[1:]:
                new_para = doc.add_paragraph()
                await format_paragraph(new_para, award, numbered=True)
                first_award_para._element.addnext(new_para._element)
                first_award_para = new_para
        else:
            default_para = doc.add_paragraph()
            await format_paragraph(default_para, "", numbered=True)
            doc.paragraphs[awards_heading_index]._element.addnext(default_para._element)
    else:
        heading_para = doc.add_paragraph("Awards and Honours")
        heading_para.style = doc.styles['Heading 1']
        
        if awards_list:
            first_award_para = doc.add_paragraph()
            await format_paragraph(first_award_para, awards_list[0], numbered=True)
            for award in awards_list[1:]:
                new_para = doc.add_paragraph()
                await format_paragraph(new_para, award, numbered=True)
                first_award_para._element.addnext(new_para._element)
                first_award_para = new_para
        else:
            default_para = doc.add_paragraph()
            await format_paragraph(default_para, "", numbered=True)
