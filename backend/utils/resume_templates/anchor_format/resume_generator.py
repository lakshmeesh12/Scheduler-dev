import json
import os
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENTATION


def set_cell_border(cell, **kwargs):
    """
    Set cell border properties
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for key, value in kwargs.items():
        if value:
            if key == 'border_bottom':
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), str(int(kwargs.get('border_thickness', 4) * 8)))  # pt to eighths of a point
                if 'border_color' in kwargs:
                    color_parts = kwargs['border_color'].lstrip('#')
                    bottom.set(qn('w:color'), color_parts)
                tcPr.append(bottom)


def set_repeat_table_header(row):
    """Set the table header to repeat on each page"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    trPr.append(tblHeader)


def rgb_from_hex(hex_color):
    """Convert hex color string to RGB tuple"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def apply_font_style(run, style_dict):
    """Apply font styles to a run object"""
    if style_dict.get('font_family'):
        run.font.name = style_dict['font_family']
   
    if style_dict.get('font_size'):
        run.font.size = Pt(style_dict['font_size'])
 
    if style_dict.get('font_color'):
        color_elem = OxmlElement('w:color')
        color_elem.set(qn('w:val'), style_dict['font_color'].lstrip('#'))
        if style_dict.get('theme_color'):
            color_elem.set(qn('w:themeColor'), style_dict['theme_color'])
        run._element.rPr.append(color_elem)
   
    run.font.bold = style_dict.get('bold', False)
    run.font.italic = style_dict.get('italic', False)
    run.font.underline = style_dict.get('underline', False)


def apply_paragraph_style(paragraph, style_dict):
    """Apply paragraph styles to a paragraph object"""
    if style_dict.get('alignment') == 'left':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif style_dict.get('alignment') == 'center':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif style_dict.get('alignment') == 'right':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    paragraph_format = paragraph.paragraph_format
    
    if 'line_spacing' in style_dict:
        if style_dict['line_spacing'] == 1.0:
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        else:
            paragraph_format.line_spacing = style_dict['line_spacing']
    
    if 'space_before' in style_dict:
        paragraph_format.space_before = Pt(style_dict['space_before'])
    
    if 'space_after' in style_dict:
        paragraph_format.space_after = Pt(style_dict['space_after'])
    
    if 'indentation' in style_dict:
        paragraph_format.left_indent = Inches(style_dict['indentation'])
    
    if 'hanging_indent' in style_dict and style_dict['hanging_indent'] > 0:
        paragraph_format.first_line_indent = -Inches(style_dict['hanging_indent'])
        paragraph_format.left_indent = Inches(style_dict['hanging_indent'])
    
    if style_dict.get('is_bullet', False):
        paragraph.style = 'List Bullet'
        paragraph_format.space_after = Pt(6)  # Adjust space after bullet points


def apply_border_bottom(paragraph, style_dict):
    """Apply bottom border to paragraph"""
    if style_dict.get('border_bottom', False):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
       
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(int(style_dict.get('border_thickness', 1) * 8)))
       
        if style_dict.get('border_color'):
            color_parts = style_dict['border_color'].lstrip('#')
            bottom.set(qn('w:color'), color_parts)
 
            if style_dict.get('border_theme_color'):
                bottom.set(qn('w:themeColor'), style_dict['border_theme_color'])
       
        pBdr.append(bottom)
        pPr.append(pBdr)


def format_date_range(start_date, end_date):
    """Format a date range for work history display"""
    if not start_date and not end_date:
        return ""
    
    if start_date:
        # Try to extract year from date string
        if isinstance(start_date, str) and len(start_date) >= 4:
            start_year = start_date[:4] if start_date[:4].isdigit() else ""
        else:
            start_year = ""
    else:
        start_year = ""
    
    if end_date:
        # Try to extract year from date string
        if isinstance(end_date, str) and len(end_date) >= 4:
            end_year = end_date[:4] if end_date[:4].isdigit() else ""
        else:
            end_year = ""
    else:
        end_year = "present"
    
    if start_year and end_year:
        return f"{start_year} to {end_year}"
    elif start_year:
        return f"{start_year} to present"
    elif end_year and end_year != "present":
        return f"Until {end_year}"
    
    return ""


def add_section_heading_with_border(document, heading_text, style_dict):
    # Add the heading
    heading_para = document.add_paragraph()
    heading_run = heading_para.add_run(heading_text)
    
    # Apply styles to the heading
    apply_font_style(heading_run, style_dict['font_style'])
    apply_paragraph_style(heading_para, style_dict['paragraph_style'])
    
    # Add bottom border
    apply_border_bottom(heading_para, style_dict['paragraph_style'])
    
    # Set special paragraph properties
    heading_para.paragraph_format.page_break_before = False
    heading_para.paragraph_format.keep_with_next = True
    
    return heading_para


def set_doc_orientation(document, payload_data, resume_styles):
    """Set the page size and document orientation to portrait or landscape"""
    sections = document.sections
    for section in sections:
        page_size = resume_styles.get('page_style', {}).get('page_size', 'Letter')
        orientation = resume_styles.get('page_style', {}).get('orientation', 'Portrait')
        
        if page_size == 'Letter':
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
        
        if orientation == 'Landscape':
            section.orientation = WD_ORIENTATION.LANDSCAPE
        
        # Set margins
        margins = resume_styles.get('page_style', {}).get('margins', {})
        section.top_margin = Inches(margins.get('top', 1.0))
        section.bottom_margin = Inches(margins.get('bottom', 1.0))
        section.left_margin = Inches(margins.get('left', 1.0))
        section.right_margin = Inches(margins.get('right', 1.0))
    return document


def add_header(document, payload_data, resume_styles):
    """Add header to the document based on the payload data and styles"""
    try: 
        if resume_styles.get('page_style', {}).get('header', False):
            # Set up header for all sections except the first one
            # We'll apply this to all sections after the first one
                for i, section in enumerate(document.sections):
                    if i > 0:  # Skip first section/page
                        header = section.header
                        
                        # First paragraph for the name
                        name_para = header.add_paragraph()
                        # credentials = payload_data.get("suffix", '')  # Add credentials logic here if needed
                        credentials = ""
                        # name_run = name_para.add_run(f"{payload_data['first_name']} {payload_data['last_name']}{credentials}")
                        name_run = name_para.add_run(payload_data['name'])

                        # Apply header name style
                        header_name_style = resume_styles['section_styles']['header_name']
                        apply_font_style(name_run, header_name_style['font_style'])
                        apply_paragraph_style(name_para, header_name_style['paragraph_style'])
                        
                        # Set right alignment for name
                        name_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        
                        # Second paragraph for the title
                        title_para = header.add_paragraph()
                        title_run = title_para.add_run(payload_data['job_title'])
                        
                        # Apply header title style
                        header_title_style = resume_styles['section_styles']['header_title']
                        apply_font_style(title_run, header_title_style['font_style'])
                        apply_paragraph_style(title_para, header_title_style['paragraph_style'])
                        
                        # Set right alignment for title
                        title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # If there aren't multiple sections yet, we need to ensure header appears starting from page 2
                # by setting different first page header
                if len(document.sections) == 1:
                    section = document.sections[0]
                    section.different_first_page_header_footer = True
                    
                    # First page header remains empty
                    
                    # Set up header for subsequent pages
                    header = section.header
                    header.is_linked_to_previous = False
                    
                    # First paragraph for the name
                    name_para = header.paragraphs[0]
                    credentials = ""  # Add credentials logic here if needed
                    # name_run = name_para.add_run(f"{payload_data['first_name']} {payload_data['last_name']}{credentials}")
                    name_run = name_para.add_run(payload_data['name'])

                    # Apply header name style
                    header_name_style = resume_styles['section_styles']['header_name']
                    apply_font_style(name_run, header_name_style['font_style'])
                    apply_paragraph_style(name_para, header_name_style['paragraph_style'])
                    
                    # Set right alignment for name
                    name_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    # Second paragraph for the title
                    title_para = header.add_paragraph()
                    title_run = title_para.add_run(payload_data['job_title'])
                    
                    # Apply header title style
                    header_title_style = resume_styles['section_styles']['header_title']
                    apply_font_style(title_run, header_title_style['font_style'])
                    apply_paragraph_style(title_para, header_title_style['paragraph_style'])
                    
                    # Set right alignment for title
                    title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        return document
    except Exception as e:
        print(f"Error adding header: {e}")
        exc_type, exc_obj, exc_tb = sys.exc_info()
        fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        message = f"{fname} : Line no {exc_tb.tb_lineno} - {exc_type} : {e}"
        raise ValueError(message)


def add_name_and_title_section(document, payload_data, resume_styles):
    """Add the name and title section to the document"""
    try:
        # Add Name and Title Section to the document
        name_style = resume_styles['section_styles']['name']
        name_para = document.add_paragraph()
        credentials = ""  # Add credentials logic here if needed
        # name_run = name_para.add_run(f"{payload_data['first_name']} {payload_data['last_name']} {credentials}")
        name_run = name_para.add_run(payload_data['name'])
        apply_font_style(name_run, name_style['font_style'])
        apply_paragraph_style(name_para, name_style['paragraph_style'])
        
        # Add external title =================================================
        title_style = resume_styles['section_styles']['title']
        title_para = document.add_paragraph()
        title_run = title_para.add_run(payload_data['job_title'])
        apply_font_style(title_run, title_style['font_style'])
        apply_paragraph_style(title_para, title_style['paragraph_style'])
        apply_border_bottom(title_para, title_style['paragraph_style'])

        return document
    except Exception as e:
        print(f"Error adding name and title section: {e}")
        exc_type, exc_obj, exc_tb = sys.exc_info()
        fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        message = f"{fname} : Line no {exc_tb.tb_lineno} - {exc_type} : {e}"
        raise ValueError(message)


def add_education(document, right_cell, payload_data, resume_styles, ecw_check):
    """Add the education section to the document"""
    try:
        edu_header_para = right_cell.paragraphs[0]
        edu_header_run = edu_header_para.add_run("Education")
        
        edu_header_style = resume_styles['section_styles']['education_header']
        apply_font_style(edu_header_run, edu_header_style['font_style'])
        apply_paragraph_style(edu_header_para, edu_header_style['paragraph_style'])
        
        # Education items
        education_items = payload_data.get('education', [])
        for edu in education_items:
            edu_item_para = right_cell.add_paragraph()
            edu_text = ""
            
            # Extract education details
            if isinstance(edu, dict):
                degree = edu.get('degree', '')
                major = edu.get('major_or_specialization', '')
                institution = edu.get('college_institution', '')
                if degree:
                    edu_text = f"{degree}, "
                if major:
                    edu_text += f"{major}, "
                if institution:
                    edu_text += f"{institution}."
                else:
                    edu_text = edu_text.rstrip(', ')

            elif isinstance(edu, str):
                edu_text = edu
                
            edu_item_run = edu_item_para.add_run(edu_text)
            
            edu_item_style = resume_styles['section_styles']['education_item']
            apply_font_style(edu_item_run, edu_item_style['font_style'])
            apply_paragraph_style(edu_item_para, edu_item_style['paragraph_style'])

        return right_cell
    except Exception as e:
        print(f"Error adding education section: {e}")
        exc_type, exc_obj, exc_tb = sys.exc_info()
        fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        message = f"{fname} : Line no {exc_tb.tb_lineno} - {exc_type} : {e}"
        raise ValueError(message)
    

def add_certifications(document, right_cell, payload_data, resume_styles, ecw_check):
    """Add the certifications section to the document"""
    try:
        # cert_header_para = right_cell.add_paragraph()
        if ecw_check == 0:
            cert_header_para = right_cell.paragraphs[0]
        else:
            cert_header_para = right_cell.add_paragraph()
        cert_header_run = cert_header_para.add_run("Licenses/Certifications")
        
        cert_header_style = resume_styles['section_styles']['certification_header']
        apply_font_style(cert_header_run, cert_header_style['font_style'])
        apply_paragraph_style(cert_header_para, cert_header_style['paragraph_style'])
    
        # Certification items
        certification_items = payload_data.get('certifications', [])
        for cert in certification_items:
            cert_item_para = right_cell.add_paragraph()
            cert_text = ""
            
            # Extract certification details
            if isinstance(cert, dict):
                cert_type = cert.get('certification', '')
                cert_state = cert.get('state', '')
                reg_no = ""

                if cert_type:
                    cert_text = f"{cert_type}, "
                if cert_state:
                    cert_text += f"{cert_state}, "
                if reg_no:
                    cert_text += f"{reg_no}."
                else:
                    cert_text = cert_text.rstrip(', ')

            elif isinstance(cert, str):
                cert_text = cert
                
            cert_item_run = cert_item_para.add_run(cert_text)
            
            cert_item_style = resume_styles['section_styles']['certification_item']
            apply_font_style(cert_item_run, cert_item_style['font_style'])
            apply_paragraph_style(cert_item_para, cert_item_style['paragraph_style'])

        return right_cell
    except Exception as e:
        print(f"Error adding certifications section: {e}")
        exc_type, exc_obj, exc_tb = sys.exc_info()
        fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        message = f"{fname} : Line no {exc_tb.tb_lineno} - {exc_type} : {e}"
        raise ValueError(message)


def add_work_history(document, right_cell, payload_data, resume_styles, ecw_check):
    """Add the work history section to the document"""
    try:
        # work_header_para = right_cell.add_paragraph()
        if ecw_check == 0:
            work_header_para = right_cell.paragraphs[0]
        else:
            work_header_para = right_cell.add_paragraph()
        work_header_run = work_header_para.add_run("Work History")
        
        work_header_style = resume_styles['section_styles']['work_history_header']
        apply_font_style(work_header_run, work_header_style['font_style'])
        apply_paragraph_style(work_header_para, work_header_style['paragraph_style'])
        
        # Iterate through work history items
        work_history_items = payload_data.get('work_history', [])
        for work in work_history_items:
            work_item_para = right_cell.add_paragraph()
            
            work_text = ""
            # Extract work history details
            company = work.get('company', '')
            designation = work.get('designation', '')
            start_year = work.get('start_year', '')
            end_year = work.get('end_year', '')
            
            # Format the work history entry
            if company:
                work_text = f"{company}, "
            if designation:
                work_text += f"{designation}, "        
            
            # Add the time period
            if start_year:
                time_period = f"{start_year} to {end_year}" if end_year else f"{start_year} to present"
                work_text += f"{time_period}"
            
            work_item_run = work_item_para.add_run(work_text)
            
            work_item_style = resume_styles['section_styles']['work_history_item']
            apply_font_style(work_item_run, work_item_style['font_style'])
            apply_paragraph_style(work_item_para, work_item_style['paragraph_style'])

        return right_cell
    except Exception as e:
        print(f"Error adding work history section: {e}")
        exc_type, exc_obj, exc_tb = sys.exc_info()
        fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        message = f"{fname} : Line no {exc_tb.tb_lineno} - {exc_type} : {e}"
        raise ValueError(message)


def add_bio_and_sidebar_section(document, payload_data, resume_styles):
    try:
        bio_table_style = resume_styles['table_styles']['bio_table']
        bio_table = document.add_table(1, 2)
        bio_table.autofit = False
        bio_table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Apply cell styles
        for i, cell in enumerate([bio_table.cell(0, 0), bio_table.cell(0, 1)]):
            cell_style = bio_table_style['cell_styles'][i]
            cell.width = Inches(bio_table_style['column_widths'][i])
            
            # Convert string vertical alignment to proper enum values
            if cell_style.get('vertical_align') == 'top':
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            elif cell_style.get('vertical_align') == 'center':
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            elif cell_style.get('vertical_align') == 'bottom':
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
            
            # Set cell margins/padding
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            
            # Top padding
            tcMar = OxmlElement('w:tcMar')
            top = OxmlElement('w:top')
            top.set(qn('w:w'), str(int(cell_style['top_padding'] * 720)))  # twips
            top.set(qn('w:type'), 'dxa')
            tcMar.append(top)
            
            # Left padding
            left = OxmlElement('w:left')
            left.set(qn('w:w'), str(int(cell_style['left_padding'] * 720)))
            left.set(qn('w:type'), 'dxa')
            tcMar.append(left)
            
            # Bottom padding
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:w'), str(int(cell_style['bottom_padding'] * 720)))
            bottom.set(qn('w:type'), 'dxa')
            tcMar.append(bottom)
            
            # Right padding
            right = OxmlElement('w:right')
            right.set(qn('w:w'), str(int(cell_style['right_padding'] * 720)))
            right.set(qn('w:type'), 'dxa')
            tcMar.append(right)
            
            tcPr.append(tcMar)
        
        # If no borders, set all borders to none
        if not bio_table_style.get('borders', True):
            for row in bio_table.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    
                    tblBorders = OxmlElement('w:tcBorders')
                    for border in ['top', 'left', 'bottom', 'right']:
                        border_elem = OxmlElement(f'w:{border}')
                        border_elem.set(qn('w:val'), 'nil')
                        tblBorders.append(border_elem)
                    
                    tcPr.append(tblBorders)

        # Add Bio content to left cell
        bio_cell = bio_table.cell(0, 0)
        bio_para = bio_cell.paragraphs[0]
        bio_text = payload_data.get("bio_sketch", "")
        bio_run = bio_para.add_run(bio_text)

        # Apply styles to the bio paragraph
        bio_style = resume_styles['section_styles']['bio']
        apply_font_style(bio_run, bio_style['font_style'])
        apply_paragraph_style(bio_para, bio_style['paragraph_style'])

        #### Add Sidebar content to right cell
        right_cell = bio_table.cell(0, 1)
        ecw_check = 0
    
        education = payload_data.get('education', [])
        if education:
            # Add Education section to the right cell
            right_cell = add_education(document, right_cell, payload_data, resume_styles, ecw_check)
            ecw_check+=1

        certifications = payload_data.get('certifications', []) 
        if certifications:
            # Add Certifications section to the right cell
            right_cell = add_certifications(document, right_cell, payload_data, resume_styles, ecw_check)
            ecw_check+=1

        work_history = payload_data.get('work_history', [])
        if work_history:
            # Add Work History section to the right cell
            right_cell = add_work_history(document, right_cell, payload_data, resume_styles, ecw_check)    
            

        return document
    except Exception as e:
        print(f"Error adding bio and sidebar section: {e}")
        exc_type, exc_obj, exc_tb = sys.exc_info()
        fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        message = f"{fname} : Line no {exc_tb.tb_lineno} - {exc_type} : {e}"
        raise ValueError(message)
    

def add_projects_section(document, payload_data, resume_styles):
    try:
        project_table_style = resume_styles['table_styles']['project_table']
        projects = payload_data.get('projects', [])
        if projects:
            # Add 1 for the header row
            project_table = document.add_table(rows=len(projects)+1, cols=1)
            project_table.autofit = False
            project_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            # Set column width
            project_table.columns[0].width = Inches(project_table_style['column_widths'][0])
            
            # Set first row as header row that repeats on new pages
            if project_table_style.get('repeat_header', False):
                set_repeat_table_header(project_table.rows[0])
            
            # Add "Relevant Project Experience" to the header row
            header_cell = project_table.cell(0, 0)
            header_para = header_cell.paragraphs[0]
            header_run = header_para.add_run("Relevant Project Experience")

            title_style = resume_styles['section_styles']['title']
            project_header_style = resume_styles['section_styles']['project_header']
            apply_font_style(header_run, project_header_style['font_style'])
            apply_paragraph_style(header_para, project_header_style['paragraph_style'])
            apply_border_bottom(header_para, title_style['paragraph_style'])
            
            # Apply border to the header row
            tc = header_cell._tc
            tcPr = tc.get_or_add_tcPr()
            
            # If no borders for other cells, set those borders to none
            if not project_table_style.get('borders', True):
                for row_idx, row in enumerate(project_table.rows):
                    # Skip the header row
                    if row_idx == 0:
                        continue
                        
                    for cell in row.cells:
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        
                        tblBorders = OxmlElement('w:tcBorders')
                        for border in ['top', 'left', 'bottom', 'right']:
                            border_elem = OxmlElement(f'w:{border}')
                            border_elem.set(qn('w:val'), 'nil')
                            tblBorders.append(border_elem)
                        
                        tcPr.append(tblBorders)
            
            # Apply cell styles
            for row_idx, row in enumerate(project_table.rows):
                for i, cell in enumerate(row.cells):
                    if i < len(project_table_style['cell_styles']):
                        cell_style = project_table_style['cell_styles'][i]
                        
                        # Fix: Convert string vertical alignment to proper enum values
                        if cell_style.get('vertical_align') == 'top':
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                        elif cell_style.get('vertical_align') == 'center':
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        elif cell_style.get('vertical_align') == 'bottom':
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
                        
                        # Set cell margins/padding
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        
                        tcMar = OxmlElement('w:tcMar')
                        
                        # Top padding
                        top = OxmlElement('w:top')
                        top.set(qn('w:w'), str(int(cell_style['top_padding'] * 720)))  # twips
                        top.set(qn('w:type'), 'dxa')
                        tcMar.append(top)
                        
                        # Left padding
                        left = OxmlElement('w:left')
                        left.set(qn('w:w'), str(int(cell_style['left_padding'] * 720)))
                        left.set(qn('w:type'), 'dxa')
                        tcMar.append(left)
                        
                        # Bottom padding
                        bottom = OxmlElement('w:bottom')
                        bottom.set(qn('w:w'), str(int(cell_style['bottom_padding'] * 720)))
                        bottom.set(qn('w:type'), 'dxa')
                        tcMar.append(bottom)
                        
                        # Right padding
                        right = OxmlElement('w:right')
                        right.set(qn('w:w'), str(int(cell_style['right_padding'] * 720)))
                        right.set(qn('w:type'), 'dxa')
                        tcMar.append(right)
                        
                        tcPr.append(tcMar)
            
            # Add projects - starting from row 1 (after the header)
            for i, project in enumerate(projects):
                # Project goes in row i+1 (skipping header row)
                cell = project_table.cell(i+1, 0)
                
                # Project title
                title_para = cell.paragraphs[0]
                
                # Get project details
                client = project.get('institution', '')
                title = project.get('title', '')
                location = project.get('location', '')

                # Format the project title text
                title_text = ""

                if title:
                    title_text += f"{title}, "
                if client:
                    title_text += f"{client}"
                else:
                    title_text = title_text[:-2]
                if location:
                    title_text += f" ({location})"

                title_text = title_text.replace("\n"," ")
                    
                title_run = title_para.add_run(title_text)
                
                project_title_style = resume_styles['section_styles']['project_title']
                apply_font_style(title_run, project_title_style['font_style'])
                apply_paragraph_style(title_para, project_title_style['paragraph_style'])
                
                # Project description
                desc_para = cell.add_paragraph()
                desc_run = desc_para.add_run(project.get('description', ''))
                
                project_desc_style = resume_styles['section_styles']['project_description']
                apply_font_style(desc_run, project_desc_style['font_style'])
                apply_paragraph_style(desc_para, project_desc_style['paragraph_style'])
        return document
    except Exception as e:
        print(f"Error adding projects section: {e}")
        exc_type, exc_obj, exc_tb = sys.exc_info()
        fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        message = f"{fname} : Line no {exc_tb.tb_lineno} - {exc_type} : {e}"
        raise ValueError(message)
    

def add_publications_section(document, payload_data, resume_styles):
    """Add the publications section to the document"""
    try:
        pub_header_style = resume_styles['section_styles']['publications_header']
        
        # Add the Publications heading
        pub_header_para = add_section_heading_with_border(document, "Publications", pub_header_style)
        
        # Add publications from payload
        publications = payload_data.get('publications', [])
        if publications:
            for i, pub in enumerate(publications):
                pub_item_para = document.add_paragraph()
                pub_item_run = pub_item_para.add_run(pub)
                
                pub_item_style = resume_styles['section_styles']['publications_item']
                apply_font_style(pub_item_run, pub_item_style['font_style'])
                apply_paragraph_style(pub_item_para, pub_item_style['paragraph_style'])
        return document
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        message = f"{fname} : Line no {exc_tb.tb_lineno} - {exc_type} : {e}"
        raise ValueError(message)


def add_other_entity(document, entity, other_entities, resume_styles):
    """Add an entity to the 'others' section in the document"""
    try:
        others_header_style = resume_styles['section_styles']['awards_header']
        item_style = resume_styles['section_styles']['awards_item']
        entity_content = other_entities.get(entity)

        if isinstance(entity_content, str) and entity_content.strip():
            add_section_heading_with_border(document, entity if entity in ["Awards and Honors"] else entity.title(), others_header_style)

            para = document.add_paragraph()
            run = para.add_run(entity_content)
            apply_font_style(run, item_style['font_style'])
            apply_paragraph_style(para, item_style['paragraph_style'])

        elif isinstance(entity_content, list) and entity_content:
            add_section_heading_with_border(document, entity if entity in ["Awards and Honors"] else entity.title(), others_header_style)

            for item in entity_content:
                para = document.add_paragraph()
                run = para.add_run(item)
                apply_font_style(run, item_style['font_style'])
                apply_paragraph_style(para, item_style['paragraph_style'])

        return document

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        message = f"{fname} : Line no {exc_tb.tb_lineno} - {exc_type.__name__} : {e}"
        raise ValueError(message)


def arrange_others_section(document, payload_data, resume_styles):
    """Arrange the 'others' section in the document"""
    try:
        other_entities = payload_data.get('other_fields', {})

        # setting priorty order for other entities
        priority_order = [
            'presentations', 'Awards and Honors'
        ]
        for entity in priority_order:
            if entity in other_entities:
                add_other_entity(document, entity, other_entities, resume_styles)
                other_entities.pop(entity)

        # Add remaining entities
        for entity in other_entities:
            if entity:
                add_other_entity(document, entity, other_entities, resume_styles)

        return document
    except Exception as e:
        print(f"Error arranging others section: {e}")
        exc_type, exc_obj, exc_tb = sys.exc_info()
        fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        message = f"{fname} : Line no {exc_tb.tb_lineno} - {exc_type} : {e}"
        raise ValueError(message)


def create_resume(payload_data, resume_styles, file_path):
    """Create a resume document based on the provided payload data and styles"""
    try:
        print(f"Creating resume with payload data and styles... \n payload_data: {payload_data.keys()} \n other_keys: {payload_data.get('other_fields', {}).keys()}")
        document = Document()

        # Set document orientation and margins
        document = set_doc_orientation(document, payload_data, resume_styles)

        # Add Header 
        document = add_header(document, payload_data, resume_styles)

        # Add Name and Title Section
        document = add_name_and_title_section(document, payload_data, resume_styles)

        # Create a Bio and Side bar section
        if "bio_sketch" in payload_data or "education" in payload_data or "certifications" in payload_data or "work_history" in payload_data:
            document = add_bio_and_sidebar_section(document, payload_data, resume_styles)

        # Add blank paragraph between tables
        spacer_para = document.add_paragraph()
        spacer_para.paragraph_format.space_before = Pt(12)
        spacer_para.paragraph_format.space_after = Pt(0)

        # Add Projects Section
        if payload_data.get('projects', []):
            document = add_projects_section(document, payload_data, resume_styles)

        # Solving Bug by adding Testimony and Presentations to other_fields
        if payload_data.get('Testimony', []):
            payload_data['other_fields'] = payload_data.get('other_fields', {})
            payload_data['other_fields']['Testimony'] = payload_data['Testimony']
            payload_data.pop('Testimony')
        
        if payload_data.get('presentations', []):
            payload_data['other_fields'] = payload_data.get('other_fields', {})
            payload_data['other_fields']['presentations'] = payload_data['presentations']
            payload_data.pop('presentations')
        
        if payload_data.get('awards', []):
            payload_data['other_fields'] = payload_data.get('other_fields', {})
            payload_data['other_fields']['awards'] = payload_data['awards']
            payload_data.pop('awards')

        with open('end_of_fun_payload.json', 'w', encoding='utf-8') as file:
            json.dump(payload_data, file, ensure_ascii=False, indent=4)

        # Add Testimony section
        if payload_data.get('other_fields', {}):
            
            # Replacing 'Testimony' with 'Litigation work' and Awards with 'Awards and Honors'
            key_mapping = {
                "Testimony": "Litigation work",
                "awards": "Awards and Honors"
            }
            payload_data['other_fields'] = {key_mapping.get(k, k): v for k, v in payload_data['other_fields'].items()}
            print(payload_data['other_fields'].keys())

            # Add Litigation work section if it exists
            if payload_data['other_fields'].get('Litigation work', []):
                document = add_other_entity(document, 'Litigation work', payload_data['other_fields'], resume_styles)
                payload_data['other_fields'].pop('Litigation work')

        # Add publications section
        if payload_data.get('publications', []):
            document = add_publications_section(document, payload_data, resume_styles)
        
        # Add other sections like awards, skills, etc.
        if payload_data.get('other_fields', {}):
            document = arrange_others_section(document, payload_data, resume_styles)

        # Saving the document
        if file_path is None:
            file_path = f"{payload_data.get('template_type', 'Resume')}-{payload_data.get('first_name', '')}{payload_data.get('last_name', '')}.docx"
        document.save(file_path)

        

        return "Resume created successfully at " + file_path

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        message = f"{fname} : Line no {exc_tb.tb_lineno} - {exc_type} : {e}"
        raise ValueError(message)


def remove_unicode_recursive(data):
    def clean_string(text):
        if isinstance(text, str):
            return text.encode('ascii', errors='ignore').decode('ascii')
        return text
   
    if isinstance(data, str):
        return clean_string(data)
   
    elif isinstance(data, dict):
        cleaned_dict = {}
        for key, value in data.items():
            clean_key = clean_string(key) if isinstance(key, str) else key
            cleaned_dict[clean_key] = remove_unicode_recursive(value)
        return cleaned_dict
   
    elif isinstance(data, list):
        return [remove_unicode_recursive(item) for item in data]
   
    elif isinstance(data, tuple):
        return tuple(remove_unicode_recursive(item) for item in data)
   
    else:
        return data


def generate_resume_from_json(payload_json, styles_json):
    """Generate a resume document from JSON data"""
    # Parse JSON data
    payload_data = json.loads(payload_json) if isinstance(payload_json, str) else payload_json
    resume_styles = json.loads(styles_json) if isinstance(styles_json, str) else styles_json

    def remove_unicode(text):
        return text.encode('ascii', errors='ignore').decode('ascii')
    
    # payload_data = remove_unicode_recursive(payload_data)

    with open('after_remove_unicodes.json', 'w', encoding='utf-8') as file:
        json.dump(payload_data, file, ensure_ascii=False, indent=4)

    if payload_data.get("bio_sketch"):
            primary = False
            for rec in payload_data['bio_sketch']:
                if rec.get('primary'):
                    payload_data['bio_sketch'] = rec.get('bio_sketch')
                    primary=True
            if primary==False:    
                payload_data['bio_sketch'] = payload_data['bio_sketch'][0].get("bio_sketch")
            if payload_data['bio_sketch']:
                payload_data['bio_sketch'] = remove_unicode(payload_data['bio_sketch'])
    if payload_data.get("publications"):
        payload_data['publications'] = list(map(lambda x: x['title'], payload_data['publications']))
    if payload_data.get('certifications'):
        payload_data['certifications'] = list(map(lambda x: f"{x['certification']}, {x.get('location',x.get('state',''))}", payload_data['certifications']))


    print(type(resume_styles), type(payload_data))
    print("Keys from Payload...",payload_data.keys())

    with open('final_payload.json', 'w', encoding='utf-8') as file:
        json.dump(payload_data, file, ensure_ascii=False, indent=4)

    # Create the document
    output_filename = f"{payload_data.get('template_type', 'Resume')}-{payload_data.get('first_name', '')}{payload_data.get('last_name', '')}.docx"
    messsage = create_resume(payload_data, resume_styles, output_filename)
    
    return output_filename


# Example usage
def main():
    with open('new_payload.json', 'r',  encoding='utf-8') as file:
        payload_json = file.read()
    with open('resume_styles.json', 'r',  encoding='utf-8') as file:
        styles_json = file.read()
    
    output_file = generate_resume_from_json(payload_json, styles_json)
    print(f"Resume generated: {output_file}")

if __name__ == "__main__":
    main()





# async def generate_resume(body: dict, User: dict = Depends(token_decoder.get_token)):
#     try:
#         file_dict = {}
#         data = body
#         user_id = data["user_id"]
#         file_dict['user_id'] = user_id
#         data.pop("user_id",None)
#         template_type = body.get('template_type', "sf330")
#         file_name = data["filename"]
#         file_dict['file_name'] = file_name
#         file_dict['file_type'] = ".docx"
#         data['name'] = f"{data['first_name']} {data['last_name']}"

#         def remove_unicode(text):
#             return text.encode('ascii', errors='ignore').decode('ascii')

#         if data.get("bio_sketch"):
#             primary = False
#             for rec in data['bio_sketch']:
#                 if rec.get('primary'):
#                     data['bio_sketch'] = rec.get('bio_sketch')
#                     primary=True
#             if primary==False:    
#                 data['bio_sketch'] = data['bio_sketch'][0].get("bio_sketch")
#             if data['bio_sketch']:
#                 data['bio_sketch'] = remove_unicode(data['bio_sketch'])
#         if data.get("publications"):
#             data['publications'] = list(map(lambda x: x['title'], data['publications']))
#         if data.get('certifications'):
#             data['certifications'] = list(map(lambda x: f"{x.get('certification', '')}, {x.get('location',x.get('state',''))}", data['certifications']))
        
#         data['other_fields'] = data.get("other_fields", {})
        
#         file_path = str(tempfile.gettempdir()).replace("\\", "/") + "/"
        
#         users = await mongo_db['users'].find_one(
#             {"user_id": user_id}, 
#             {'experience_start_date': 1, "hire_date": 1}
#         )

#         start_date = users['experience_start_date']
#         start_date = start_date
#         hire_date = users['hire_date']
#         current_exp = int((datetime.now() - hire_date).days//365.25)
#         total_experience = int((datetime.now() - start_date).days//365.25)

#         data['total_experience'] = total_experience
#         data['years_experience_firm'] = current_exp

#         # template = ResumeGenerationTemplate(data)
#         # html_content = await run_chatgpt(template.user_prompt, template.system_prompt, 0.6)
#         # html_content = html_content.replace('```','').replace('markdown', '')
#         # logger.info("html content",html_content)
#         # output_path = await convert_markdown_to_docx(html_content, file_path+file_name)
#         file_name = file_name + ".docx"
#         output_path = file_path+file_name
#         if template_type == "AnchorQea":
#             # await initialize_config()
#             # output_path = await resume_generation_anchor(
#             #     data, 
#             #     CURRENT_DIR+"/utils/resume_templates/anchor_format/AnchorTemplate.docx", 
#             #     file_path+file_name
#             # )
#             template_dir = CURRENT_DIR + "/utils/resume_templates/anchor_format/"
#             with open(template_dir+'resume_styles.json', 'r',  encoding='utf-8') as file:
#                 styles_json = json.load(file)

#             create_resume(
#                 data,
#                 styles_json,
#                 file_path+file_name
#             )

#         elif template_type == "sf330":
#             output_path = await resume_generation_sf330(
#                 data, 
#                 CURRENT_DIR+"/utils/resume_templates/sf330_format/Template form sf330.docx", 
#                 file_path+file_name
#             )
#         logger.info(output_path)
       
#         s3_path = await upload_file_to_s3(file_path+file_name,"generate_resume", user_id,)
#         if s3_path:
#             file_dict['file_path'] = s3_path
#             file_dict['resume_id'] = str(uuid4())
#             file_dict['process_status'] = 'Completed'
#             file_dict['internal_process_status'] = "Generated"
#             file_dict['last_modified_date'] = str(datetime.now(timezone.utc))
#         if os.path.exists(output_path):
#             os.remove(output_path)
#             logger.info("Temporary file deleted successfully")
       
#         await mongo_db['generatedresume'].update_one(
#             {"user_id": user_id, "file_name": file_name},
#             {"$set": file_dict},
#             upsert=True
#         )
#         logger.info("Inserted record in mongodb!")
#         return JSONResponse(
#             content={
#                 "message": "Successfully generated resume!",
#                 "body": s3_path
#             },
#         status_code=200)

#     except Exception as err:
#         exc_type, exc_obj, exc_tb = sys.exc_info()
#         fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
#         message = f"{fname} : Line no {exc_tb.tb_lineno} - {exc_type} : {err}"
#         logger.info(f"Error while generating resume: {message}")
#         # message = f"Error while generating resume: {err} line number {line_number}"
#         logger.error(message)
#         return JSONResponse(
#             content={
#                 "message": message
#             },
#             status_code=500)
 
