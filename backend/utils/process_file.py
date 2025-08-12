import os
import re
from io import BytesIO
from typing import Dict
from fastapi import HTTPException
from fpdf import FPDF
from docx import Document
from datetime import datetime
from db.mongo.config import db as mongo_db
from bs4 import BeautifulSoup
import markdown
import aiofiles
import asyncio
from typing import Optional
import tempfile


async def convert_markdown_to_docx(markdown_content: str, output_path: Optional[str] = None) -> str:
    try:
        # Convert markdown to HTML
        html = markdown.markdown(markdown_content)
        soup = BeautifulSoup(html, 'html.parser')
        
        # Create document
        doc = Document()
        
        # Process content
        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'code']):
            if element.name.startswith('h'):
                level = int(element.name[1])
                doc.add_heading(element.get_text(), level)
            elif element.name in ['ul', 'ol']:
                for item in element.find_all('li'):
                    doc.add_paragraph(item.get_text(), style='List Bullet')
            elif element.name == 'code':
                doc.add_paragraph(element.get_text(), style='No Spacing')
            else:
                doc.add_paragraph(element.get_text())

        if not output_path:
            temp_dir = tempfile.gettempdir()
            output_path = os.path.join(temp_dir, f'converted_{asyncio.current_task().get_name()}.docx')

        await asyncio.to_thread(doc.save, output_path)
        return output_path

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")

async def convert_html_to_docx(html_content: str, output_path: Optional[str] = None) -> str:
    """
    Asynchronously converts HTML content to DOCX format.
    
    Args:
        html_content: HTML string to convert
        output_path: Optional path for output file
        
    Returns:
        Path to the created DOCX file
    """
    try:
        soup = BeautifulSoup(html_content, 'html.parser')
        # Create document
        doc = Document()
        
        # Process text content
        for paragraph in soup.find_all(['p', 'div', 'h1', 'h2', 'h3']):
            if paragraph.name.startswith('h'):
                level = int(paragraph.name[1])
                doc.add_heading(paragraph.get_text(), level)
            else:
                doc.add_paragraph(paragraph.get_text())

        # Create temporary file if no output path provided
        if not output_path:
            temp_dir = tempfile.gettempdir()
            output_path = os.path.join(temp_dir, f'converted_{asyncio.current_task().get_name()}.docx')

        # Save document
        await asyncio.to_thread(doc.save, output_path)
        return output_path

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")

def write_to_docx(file_content, file_path):
    print("****************",file_path)
    # Decode binary content if necessary
    if isinstance(file_content, bytes):
        print("stage 1")
        file_content = BytesIO(file_content).decode('utf-8', 'ignore')
        file_content = ''.join(ch for ch in file_content if ch.isprintable())
    print("stage 2")    
    document = Document()
    document.add_paragraph(file_content)
    print("stage 3")
    document.save(file_path)
    print("stage 4")
    print(f"DOCX file written to: {file_path}")

class PDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=15)
        self.add_page()
        self.set_font("Arial", size=12)

    def write_text(self, text):
        self.multi_cell(0, 10, text)

def write_to_pdf(file_content, file_path):
    # Decode binary content if necessary
    if isinstance(file_content, bytes):
        file_content = file_content.decode("utf-8", "ignore")

def extract_data(resume_text):
    """Extracts data from a Markdown resume.

    Args:
        resume_text: The text content of the resume in Markdown format.

    Returns:
        A dictionary containing extracted data.
    """

    data = {
        "education": [],
        "experience": [],
        "skills": [],
        "links": []
    }

    # Extract links
    links = re.findall(r'\[(.*?)\]\((.*?)\)', resume_text)
    for link_text, link_url in links:
        data['links'].append({'text': link_text, 'url': link_url})

    # Identify sections using keywords
    sections = ['Education', 'Experience', 'Skills']
    current_section = None
    for line in resume_text.split('\n'):
        for section in sections:
            if section in line:
                current_section = section
                break

        if current_section:
            if current_section == 'Education':
                # Extract education details (adjust regex as needed)
                match = re.search(r'^\*\s*(.*?)\s*\*\s*(.*?)\s*\((.*?)\)', line)
                if match:
                    degree, institution, years = match.groups()
                    data['education'].append({'degree': degree, 'institution': institution, 'years': years})
            elif current_section == 'Experience':
                # Extract experience details (adjust regex as needed)
                match = re.search(r'^\*\s*(.*?)\s*\*\s*(.*?)\s*\((.*?)\)', line)
                if match:
                    title, company, years = match.groups()
                    data['experience'].append({'title': title, 'company': company, 'years': years})
            elif current_section == 'Skills':
                # Extract skills (adjust regex as needed)
                skills = line.strip().split(', ')
                data['skills'].extend(skills)

    return data

async def separate_fields(model):
    # Helper function to calculate duration (years)
    def calculate_duration(start_date, end_date):
        if start_date and end_date:
            delta = end_date - start_date
            return delta.days / 365.25
        return 0.0

    fields = {'user': [], 
              'work_history': [], 
              'education': [], 
              'certifications': [],
              'projects': [],
              'skills': [],
              'publications': []
            }
    
    fields['certifications'] = model.get("certifications")
    
    fields['user'] = dict({
        "name": model.get("name"),
        "location": model.get("locations")
    })
    
    rec = await mongo_db['users'].find_one({"name": model['name']})
    user_id = rec['user_id']
    if not rec:
        await mongo_db['Users'].insert_one(fields['user'])
        rec = await mongo_db['users'].find_one({"name": model['name']})
        user_id = rec.user_id
    
    
    # Populate "work_history" collection
    for work in model.get("work_history", []):
        experience = {
            "company": work.get("company"),
            "designation": work.get("title"),
            "start_date": datetime.strptime(work.get("start_date", "1900-01-01"), "%Y-%m-%d"),
            "end_date": datetime.strptime(work.get("end_date", "1900-01-01"), "%Y-%m-%d"),
            "duration": calculate_duration(
                datetime.strptime(work.get("start_date", "1900-01-01"), "%Y-%m-%d"),
                datetime.strptime(work.get("end_date", "1900-01-01"), "%Y-%m-%d")
            )
        }
        await mongo_db['work_history'].insert_one(
            {
                "user_id": user_id,
                "user_name": model['name'],
                
            }
        )
        fields["work_history"].append(experience)

    # Populate "education" collection
    for edu in model.get("education", []):
        education = {
            "institute": edu.get("institution"),
            "degree": edu.get("degree"),
            "field": edu.get("field"),
            "grades": None,  # No grades provided in the model, set to None
            "year": edu.get("start_year"),
            "end_year": edu.get("end_year"),
            "duration": edu.get("duration")
        }
        fields["education"].append(education)

    # Convert "projects" field to the required schema
    projects = []
    for project in model.get("projects", []):
        proj = {
            "title": project.get("title"),
            "description": project.get("description"),
            "start_date": datetime.strptime(project.get("from_date", "1900-01-01"), "%Y-%m-%d"),
            "end_date": datetime.strptime(project.get("to_date", "1900-01-01"), "%Y-%m-%d"),
            "duration": calculate_duration(
                datetime.strptime(project.get("from_date", "1900-01-01"), "%Y-%m-%d"),
                datetime.strptime(project.get("to_date", "1900-01-01"), "%Y-%m-%d")
            )
        }
        projects.append(proj)
    fields["projects"] = projects

    # Return final fields collection
    return fields
    
async def process_filters(filters, field_map, exp="", emails=[]):
    try:
        proficiency_map = {
            "Novice": 1,
            "Intermediate": 2,
            "Advanced": 3,
            "Expert": 4
        }
        cert_mails = []

        rec = filters.get("location")
        if isinstance(rec, list):
           exp+=f" AND (location in {rec})"
        elif type(rec)==str and rec:
            exp+=f"AND (location=='{rec}')"
        
        rec = filters.get("licensesandcertification",[])
        lxp = ""
        if isinstance(rec, list) and rec:
            if not exp:
                # rec = field_map["certifications"](rec)
                for r in rec:
                    license, fl = r['license'], r['locations']
                    license = re.sub(r'[^A-Za-z]+', ' ', license).strip()
                    r = f"(content like '%{license}%')"
                    if fl:
                        r +=" AND " + "(" + " OR ".join([f"content like '%{x}%'" for x in fl]) + ")"
                    lxp += " OR " + r
                lxp = lxp.lstrip(" OR ")
                exp+= f" AND (type=='certifications' AND ({lxp}))"
            else:
                projection = {"user_id": 1}

                certs = []
                for r in rec:
                    license, locations = r['license'], r.get('locations', '')
                    license = re.escape(license)
                    pattern = f'"certification":"{license}"'
                    if locations:
                        for loc in  locations:
                            loc = re.escape(loc)
                            pattern = f'"certification":"{license}","state":"{loc}"'
                            certs.append({"certifications_json": {"$regex": pattern}})
                    else:
                        certs.append({"certifications_json": {"$regex": pattern}})

                qry = {"$or": certs}

                async for record in mongo_db['users'].find(qry, projection):
                    cert_mails.append(record['user_id'])

                # exp+= f" AND (user_id in {cert_mails})"
                # print(exp)

        elif type(rec)==str and rec:
            rec = f"(type == 'certifications' AND content like '%{rec}%')"
            rec = rec.lstrip(" OR ")
            exp+=f" AND (({exp}) AND ({rec}))"
        
        rec = filters.get("job_title")
        if type(rec)==str and rec:
            rec = [rec]
        if isinstance(rec, list) and rec:
            rec = [""]+rec
            rec = field_map["job_title"](rec)
            rec = rec.strip(" OR ")
            exp+=f" AND ({rec})"
        
        min_exp = 0
        max_exp = 80  
        if filters.get('experience'):
            exp_ranges = filters.get('experience')
            min_list = []
            max_list = []
            for ex in exp_ranges:
                numbers = list(map(int, re.findall(r'\d+', ex)))
                if len(numbers) == 2:
                    min_list.append(numbers[0])
                    max_list.append(numbers[1])
                elif len(numbers) == 1:
                    min_list.append(numbers[0])
                    max_list.append(80) 
        
            if min_list:
                min_exp = min(min_list)
            if max_list:
                max_exp = max(max_list)
            print(min_exp, max_exp)
        
        exp = exp.strip(" AND ")
        exp = exp.strip(" OR ")
        print("filter expression", exp)
        return (exp, cert_mails, min_exp, max_exp)
    except Exception as err:
        print(err)
        min_exp = 0,
        max_exp = 80
        print("filter expression", exp)
        return (exp, [], min_exp, max_exp)


async def userskills_pipeline(skill_requirements, email_ids=None):
    """
    Build an aggregation pipeline for MongoDB to retrieve skill records
    for users matching skill requirements and optional email filters.
    
    Returns fields: email, user_id, skills, skillsCategory, proficiencyLevel
    """
    # Match conditions based on category and skills
    match_conditions = [
        {
            "skillsCategory": req["skillcategory"],
            "skills": { "$in": req["skills"] }
        }
        for req in skill_requirements
    ]

    # Base match stage with $or on skill conditions
    match_stage = {
        "$match": {
            "$and": [
                { "$or": match_conditions }
            ]
        }
    }

    # Optional email filter
    if email_ids:
        match_stage["$match"]["$and"].append({
            "email": { "$in": email_ids }
        })

    # Project only the needed fields — no _id, include user_id
    project_stage = {
        "$project": {
            "_id": 0,
            "user_id": 1,
            "email": 1,
            "skills": 1,
            "skillsCategory": 1,
            "proficiencyLevel": 1
        }
    }

    return [match_stage, project_stage]


PROJECT_SECTION_HEADINGS = [
    "Project Experience", "Relevant Project Experience", "Projects",
    "Technical Projects", "Professional Projects", "Major Projects",
    "Selected Projects", "Project Highlights",

    # work
    "Work History", "Professional Experience", "Employment History", "Work Experience",
    "Career Summary", "Professional Background", "Relevant Experience",
    "Employment Experience", "Experience",
]

OTHER_SECTION_HEADINGS = [
    # education
    "Education", "Academic Background", "Qualifications", "Degrees",
    "Educational Background", "Academic Qualifications", "Education and Training",
    "Academic History", "Educational History", "Scholastic Record",

    # publications
    "Publications", "Other Publications", "Research Publications",
    "Published Work", "List of Publications", "Scientific Publications",
    "Academic Publications", "Articles and Publications",

    # presentations
    "Presentations", "Conference Presentations", "Talks",
    "Invited Talks", "Academic Presentations", "Lectures", "Seminars",

    # awards
    "Awards", "Honors", "Achievements", "Recognitions", "Distinctions",
    "Honours and Awards", "Fellowships and Awards", "Scholarships", "Prizes",

    # licenses
    "Licenses", "Certifications", "Licenses/Certifications", "Professional Licenses",
    "Credentials", "Accreditations", "Certifications and Licenses", "Certifications & Licenses",
    "Licensure", "Certs", "Certificates", "License & Certification", "License/Certification",

    # others
    "Organizations/Affiliations", "Memberships", "Affiliations",
    "Skills", "Technical Skills", "Professional Affiliations",
    "Languages", "References", "Volunteer Experience", "Training",
    "Leadership", "Conferences", "Summary", "Objective", "Interests"
]

async def extract_project_section(text):
    text_between_paragraph  = re.findall(r'>(.*?)<', text, re.DOTALL)
    text = '\n'.join(text_between_paragraph)

    # Build regex patterns (case insensitive, \b word boundaries)
    project_pattern = r'\n\n(' + '|'.join(re.escape(h) for h in PROJECT_SECTION_HEADINGS) + r')\b'
    other_pattern = r'\n\n(' + '|'.join(re.escape(h) for h in OTHER_SECTION_HEADINGS) + r')\b'

    # Search for first project section
    project_match = re.search(project_pattern, text, re.IGNORECASE)
    if not project_match:
        print("No projects available")
        return text

    # Find position after project section heading
    start_index = project_match.end()

    # Search for next other section after project section
    other_match = re.search(other_pattern, text[start_index:], re.IGNORECASE)
    if other_match:
        end_index = start_index + other_match.start()
        project_content = text[start_index:end_index]
    else:
        # If no other section found, take till end
        project_content = text[start_index:]

    print("Extracted Project Section:")
    return project_content.strip()

async def compute_duration(records):
    try:
        for rec in records:
            duration = None
            from_date = rec.get("from_date")
            to_date = rec.get('to_date')
            if from_date and to_date:
                try:
                    from_date = int(from_date.split("-")[0])
                    to_date = datetime.now().year if to_date.lower()=='present' or to_date.lower()=='ongoing' else int(to_date.split("-")[0])
                    duration = int(to_date) - from_date
                    duration = None if duration==0 else duration
                except:
                    duration = None
            # rec['from_date'] = str(from_date) if from_date else None
            # rec['to_date'] = str(to_date) if to_date else None
            rec['duration'] = duration
        return records
    except Exception as err:
        print("Error while computing duration: ",err)
        return records

if __name__ == "__main__":
    pass